from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(0.9)
section.right_margin  = Inches(0.9)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)

# ── Colours ───────────────────────────────────────────────────────────────────
PURPLE   = RGBColor(0x6B, 0x21, 0xA8)
DARK     = RGBColor(0x1F, 0x2A, 0x3C)
GREY     = RGBColor(0x6B, 0x72, 0x80)
GREEN    = RGBColor(0x06, 0x6B, 0x2B)
AMBER    = RGBColor(0x92, 0x40, 0x0E)
RED      = RGBColor(0x99, 0x17, 0x17)
BLUE     = RGBColor(0x1E, 0x40, 0xAF)
TEAL     = RGBColor(0x0F, 0x76, 0x6E)
CODE_BG  = 'F1F5F9'
HDR_BG   = '6B21A8'
ALT_BG   = 'F5F3FF'

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_para_border(para, color='6B21A8', sz='4'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('bottom',):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '1');    b.set(qn('w:color'), color)
        pBdr.append(b)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, HDR_BG)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9.5)
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = ALT_BG if ri % 2 == 1 else 'FFFFFF'
        for ci, txt in enumerate(row_data):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(9.5)
            col = DARK
            for kw, clr in [('DONE',GREEN),('PARTIAL',AMBER),('MISSING',RED),
                             ('HIGH',RED),('MEDIUM',AMBER),('LOW',GREEN),
                             ('COMPLETE',GREEN)]:
                if kw in str(txt): col = clr; break
            r.font.color.rgb = col
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t

def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb = PURPLE
        r.font.bold = True
        r.font.size = Pt(18)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb = PURPLE
        r.font.bold = True
        r.font.size = Pt(14)

def h3(text, colour=DARK):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = colour
        r.font.bold = True
        r.font.size = Pt(12)

def para(text, bold=False, colour=DARK, size=10.5, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Inches(indent * 0.22)
    r = p.add_run(text)
    r.font.bold = bold; r.font.color.rgb = colour; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, colour=DARK, size=10, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent * 0.22)
    r = p.add_run(text)
    r.font.color.rgb = colour; r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)

def code_block(lines):
    """Render a code block as a table cell."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    set_cell_bg(c, CODE_BG)
    p = c.paragraphs[0]
    run = p.add_run(lines)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    set_para_border(p)
    p.paragraph_format.space_after = Pt(6)

def label_value(label, value, label_colour=PURPLE, val_colour=DARK):
    p = doc.add_paragraph()
    r1 = p.add_run(label + '  ')
    r1.font.bold = True; r1.font.color.rgb = label_colour; r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.color.rgb = val_colour; r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

def tag(text, colour):
    """Inline coloured tag."""
    return f'[{text}]'

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('BillingTool')
r.font.bold = True; r.font.size = Pt(36); r.font.color.rgb = PURPLE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Gap Analysis — Detailed Findings & Proposed Solutions')
r2.font.size = Pt(16); r2.font.color.rgb = DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Generated: 2026-04-23')
r3.font.size = Pt(11); r3.font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Introduction')
divider()
para(
    'This document provides a deep-dive technical analysis of every identified gap in the '
    'BillingTool application. For each gap the following is documented: the exact current '
    'state of the code (with file paths and line numbers), the root cause of the gap, a '
    'concrete proposed solution with code examples, and an estimated implementation effort. '
    'Gaps are ordered by business impact.'
)
para(
    'Stack: React 18 + TypeScript (Vite) frontend · CodeIgniter 4 PHP 8.1 backend · '
    'MySQL · jsPDF for client-side PDF · Zustand for state management.'
)
doc.add_paragraph()

# Summary table
h2('1.1  Summary of All Gaps')
rows = [
    ('1',  'Template Designer drag-drop',       'HIGH',   'PARTIAL',  '3–4 h'),
    ('2',  'Invoice status transition enforcement','HIGH', 'MISSING',  '2 h'),
    ('3',  'Company logo upload in Settings',   'HIGH',   'PARTIAL',  '1–1.5 h'),
    ('4',  'Invoice number format builder UI',  'HIGH',   'PARTIAL',  '2–2.5 h'),
    ('5',  'Admin RBAC',                        'HIGH',   'MISSING',  '3–4 h'),
    ('6',  'Two-Factor Authentication (2FA)',   'MEDIUM', 'MISSING',  '5–6 h'),
    ('7',  'Invoice locking after "sent"',      'MEDIUM', 'MISSING',  '2 h'),
    ('8',  'Email notifications',               'MEDIUM', 'MISSING',  '4–5 h'),
    ('9',  'Batch PDF export',                  'MEDIUM', 'PARTIAL',  '3–4 h'),
    ('10', 'Letter body rich-text editor',      'MEDIUM', 'PARTIAL',  '1.5–2 h'),
    ('11', 'customerApi.ts interceptor',        'LOW',    'PARTIAL',  '1–1.5 h'),
    ('12', 'Wiki admin edit UI',                'LOW',    'MISSING',  '3–4 h'),
    ('13', 'PDF/A-3 with embedded UBL',         'LOW',    'MISSING',  '6–8 h'),
    ('14', 'Invoice clone (persist to DB)',     'LOW',    'PARTIAL',  '1–1.5 h'),
    ('15', 'AuditLog field-level tracking',     'LOW',    'PARTIAL',  '2–3 h'),
]
add_table(
    ['#', 'Gap', 'Priority', 'Status', 'Effort'],
    rows,
    col_widths=[0.25, 2.85, 0.80, 0.85, 0.75]
)
para('Total estimated effort: 45–60 hours of focused development.', colour=GREY, size=9.5)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  GAPS 1-15
# ══════════════════════════════════════════════════════════════════════════════

# ─────────────────────────────────────────────────────────────────────────────
# GAP 1
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 1 — Template Designer Drag-Drop (HIGH)')
divider()

label_value('Status:', 'PARTIAL — pointer events 70% wired, but no visual feedback or smooth constraints', val_colour=AMBER)
label_value('File:', 'src/components/invoice/TemplateDesignLayout.tsx (lines 26–110, 316–322)')
label_value('Effort:', '3–4 hours')

h2('Current State')
para(
    'The state variables exist (draggingId, dragStartPos, elementStartPos, selectedId, '
    'snapToGrid, zoom — lines 26–35) and native pointer event handlers are attached '
    '(handlePointerDown lines 75–84, handlePointerMove lines 86–106, handlePointerUp '
    'line 108–110). However three critical pieces are missing:'
)
for item in [
    'No visual feedback while dragging — no cursor change, no opacity, no ghost element',
    'No smooth boundary constraint enforcement during the drag motion (only applied on release)',
    'Elements of type "items", "notes", "footer" are excluded from drag (line 76) with no visible indicator',
    'No drag library (react-dnd / dnd-kit) — relies entirely on native pointer events which lack cross-device consistency',
]:
    bullet(item, colour=AMBER)

h2('Root Cause')
para(
    'The drag infrastructure was scaffolded but visual polish was never added. The pointer '
    'handlers update position state correctly, but the rendered elements do not reflect '
    'the in-flight drag state (no CSS transform, no cursor: grabbing, no opacity: 0.7).'
)

h2('Proposed Solution')
para('Step 1 — Add drag visual feedback via CSS class toggling:', bold=True)
code_block(
"""// TemplateDesignLayout.tsx — enhance handlePointerMove (lines 86–106)
const handlePointerMove = (e: React.PointerEvent) => {
  if (!draggingId || !dragStartPos || !elementStartPos) return;

  const dx = (e.clientX - dragStartPos.x) / zoom;
  const dy = (e.clientY - dragStartPos.y) / zoom;

  let newX = elementStartPos.x + dx;
  let newY = elementStartPos.y + dy;

  // Snap to grid
  if (snapToGrid) {
    newX = Math.round(newX / GRID_SIZE) * GRID_SIZE;
    newY = Math.round(newY / GRID_SIZE) * GRID_SIZE;
  }

  // Boundary constraints (enforced during move, not only on release)
  newX = Math.max(0, Math.min(newX, A4_WIDTH_PT - 40));
  newY = Math.max(0, Math.min(newY, A4_HEIGHT_PT - 20));

  setLayout(prev => prev.map(el =>
    el.id === draggingId ? { ...el, x: newX, y: newY } : el
  ));
};

// In element JSX (lines 316–322), add dynamic styles:
<div
  key={el.id}
  onPointerDown={canDrag ? (e) => handlePointerDown(e, el.id) : undefined}
  style={{
    left: el.x, top: el.y, width: el.width,
    cursor: draggingId === el.id ? 'grabbing' : (canDrag ? 'grab' : 'default'),
    opacity: draggingId === el.id ? 0.75 : 1,
    outline: selectedId === el.id ? '2px solid #6B21A8' : 'none',
    transition: draggingId === el.id ? 'none' : 'outline 0.1s',
    userSelect: 'none',
    zIndex: draggingId === el.id ? 999 : el.zIndex ?? 1,
  }}
>
"""
)

para('Step 2 — Add "non-draggable" visual hint:', bold=True)
code_block(
"""// Show lock icon on non-draggable elements
{!canDrag && (
  <div className="absolute top-0 right-0 p-0.5 opacity-40">
    <Lock className="h-3 w-3 text-slate-400" />
  </div>
)}
"""
)

para('Step 3 — Consider migrating to @dnd-kit/core for production-grade drag-drop:', bold=True)
code_block(
"""// Install:  npm install @dnd-kit/core @dnd-kit/utilities
// Wrap canvas in <DndContext>; each element becomes a <Draggable>.
// Provides accessibility (keyboard drag), touch support, collision detection.
"""
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 2
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 2 — Invoice Status Transition Enforcement (HIGH)')
divider()

label_value('Status:', 'MISSING — any status can be set to any other in the backend', val_colour=RED)
label_value('File:', 'api/app/Controllers/InvoiceController.php (update method, lines 219–258)')
label_value('Effort:', '2 hours')

h2('Current State')
para(
    'The update() method accepts a new status value from the request body and sets it '
    'directly on the database record. Lines 247–248 only determine which log message to '
    'write — they do not validate whether the transition is allowed. An invoice with status '
    '"paid" can be set back to "draft" by a single API call.'
)
code_block(
"""// Current code (line ~247):
if ($dbData['status'] === 'validated') $action = 'validated';
if ($dbData['status'] === 'sent')      $action = 'sent';
// Direct DB update with no prior state check
$model->update($id, $dbData);
"""
)

h2('Valid Transition Matrix')
add_table(
    ['From Status', 'Allowed Transitions'],
    [
        ('draft',     'validated, deleted'),
        ('validated', 'sent, draft (revert), cancelled'),
        ('sent',      'paid, cancelled'),
        ('paid',      '(terminal — no transitions)'),
        ('cancelled', '(terminal — no transitions)'),
    ],
    col_widths=[1.5, 4.5]
)

h2('Proposed Solution')
para('Add a private validation method and call it before the DB update:', bold=True)
code_block(
"""// api/app/Controllers/InvoiceController.php

private function validateStatusTransition(string $from, string $to): void
{
    $allowed = [
        'draft'     => ['validated', 'deleted'],
        'validated' => ['sent', 'draft', 'cancelled'],
        'sent'      => ['paid', 'cancelled'],
        'paid'      => [],
        'cancelled' => [],
    ];

    if (!array_key_exists($from, $allowed) ||
        !in_array($to, $allowed[$from], true))
    {
        throw new \\InvalidArgumentException(
            "Invalid status transition: {$from} → {$to}"
        );
    }
}

// In update() method, before $model->update():
$existing = $model->find($id);
if (isset($dbData['status']) && $dbData['status'] !== $existing['status']) {
    try {
        $this->validateStatusTransition($existing['status'], $dbData['status']);
    } catch (\\InvalidArgumentException $e) {
        return $this->fail($e->getMessage(), 422);
    }
}
"""
)

para('Frontend should also grey out invalid status options in the status dropdown:', bold=True)
code_block(
"""// InvoiceEditor.tsx or InvoiceList.tsx
const ALLOWED_TRANSITIONS: Record<string, string[]> = {
  draft:     ['validated'],
  validated: ['sent', 'draft', 'cancelled'],
  sent:      ['paid', 'cancelled'],
  paid:      [],
  cancelled: [],
};

// In status select:
{STATUS_OPTIONS.map(opt => (
  <SelectItem
    key={opt.value}
    value={opt.value}
    disabled={!ALLOWED_TRANSITIONS[invoice.status]?.includes(opt.value)}
  >
    {opt.label}
  </SelectItem>
))}
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 3
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 3 — Company Logo Upload in Settings (HIGH)')
divider()

label_value('Status:', 'PARTIAL — Settings.tsx has a text URL input only; file upload is in TemplateEditor but not reused', val_colour=AMBER)
label_value('Files:', 'src/components/screens/Settings.tsx (lines 287–319)  |  src/components/invoice/TemplateEditor.tsx (lines 106–141)')
label_value('Effort:', '1–1.5 hours')

h2('Current State')
para(
    'Settings.tsx (line 291–299) renders a plain text <Input> that accepts a URL string '
    'for logoUrl. TemplateEditor.tsx (lines 106–141) already has a complete file-upload '
    'implementation: file type validation, 2 MB size limit, and FileReader base64 conversion. '
    'The same logic just needs to be ported to Settings.tsx.'
)
code_block(
"""// Settings.tsx — current (line 291):
<Input
  value={editedProfile.logoUrl || ''}
  onChange={e => handleChange('logoUrl', e.target.value)}
  placeholder="https://example.com/logo.png"
/>
// → No file picker, user must type/paste a URL manually
"""
)

h2('Proposed Solution')
code_block(
"""// Settings.tsx — add these near the top of the component:
import { Upload, X } from 'lucide-react';
const logoInputRef = useRef<HTMLInputElement>(null);

const handleLogoUpload = (e: React.ChangeEvent<HTMLInputElement>) => {
  const file = e.target.files?.[0];
  if (!file) return;

  if (!file.type.startsWith('image/')) {
    toast.error(t('settings.logoInvalidType') || 'Please upload an image file');
    return;
  }
  if (file.size > 2 * 1024 * 1024) {
    toast.error(t('settings.logoTooLarge') || 'Image must be less than 2 MB');
    return;
  }

  const reader = new FileReader();
  reader.onloadend = () => {
    handleChange('logoUrl', reader.result as string);
    toast.success(t('settings.logoUploaded') || 'Logo uploaded');
  };
  reader.readAsDataURL(file);
};

// Replace the URL <Input> block (line 291–299) with:
<div className="space-y-2">
  <Label>{t('settings.companyLogo') || 'Company Logo'}</Label>

  <input
    ref={logoInputRef}
    type="file"
    accept="image/*"
    onChange={handleLogoUpload}
    className="hidden"
  />

  <div className="flex items-center gap-3">
    {editedProfile.logoUrl ? (
      <div className="relative">
        <img
          src={editedProfile.logoUrl}
          alt="Logo preview"
          className="h-16 w-auto rounded border border-slate-200 object-contain"
        />
        <button
          type="button"
          onClick={() => handleChange('logoUrl', '')}
          className="absolute -top-2 -right-2 bg-red-500 text-white rounded-full p-0.5"
        >
          <X className="h-3 w-3" />
        </button>
      </div>
    ) : (
      <div className="h-16 w-32 border-2 border-dashed border-slate-300 rounded flex items-center justify-center text-slate-400 text-xs">
        No logo
      </div>
    )}

    <Button
      type="button"
      variant="outline"
      size="sm"
      onClick={() => logoInputRef.current?.click()}
    >
      <Upload className="h-4 w-4 mr-2" />
      {editedProfile.logoUrl ? 'Replace Logo' : 'Upload Logo'}
    </Button>
  </div>
</div>
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 4
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 4 — Invoice Number Format Builder UI (HIGH)')
divider()

label_value('Status:', 'PARTIAL — DB field and type exist; Settings.tsx shows only a plain text input with no preview or validation', val_colour=AMBER)
label_value('Files:', 'src/components/screens/Settings.tsx (lines 405–418)  |  src/types/invoice.ts (line 164)')
label_value('Effort:', '2–2.5 hours')

h2('Current State')
para(
    'The CompanyProfile interface (invoice.ts:164) exposes invoiceNumberFormat as an optional '
    'string. Settings.tsx line 405–418 renders a single <Input> with placeholder '
    '"INV-{YYYY}-{NNNNN}". There is no live preview, no variable reference guide, and '
    'no quick-insert buttons. Users must know the token syntax from memory.'
)

h2('Supported Token Reference')
add_table(
    ['Token', 'Expands To', 'Example'],
    [
        ('{YYYY}', 'Full 4-digit year',  '2026'),
        ('{YY}',   '2-digit year',       '26'),
        ('{MM}',   '2-digit month',      '04'),
        ('{N}',    '1-digit sequence',   '7'),
        ('{NNN}',  '3-digit zero-padded','007'),
        ('{NNNNN}','5-digit zero-padded','00007'),
        ('{TENANT}','Company short code', 'ACME'),
    ],
    col_widths=[1.0, 2.2, 1.4]
)

h2('Proposed Solution')
code_block(
"""// Settings.tsx — replace plain input with format builder block

// Helper to generate a live preview
const previewFormat = (fmt: string, counter = 7): string => {
  const now = new Date();
  return fmt
    .replace('{YYYY}', String(now.getFullYear()))
    .replace('{YY}',   String(now.getFullYear()).slice(-2))
    .replace('{MM}',   String(now.getMonth() + 1).padStart(2, '0'))
    .replace(/{N+}/g,  (m) => String(counter).padStart(m.length - 2, '0'))
    .replace('{TENANT}', profile.name?.slice(0, 4).toUpperCase() || 'ACME');
};

// JSX — replace the <Input> block at line 405:
<div className="space-y-3">
  <Label>{t('settings.invoiceNumberFormat')}</Label>

  <div className="flex gap-2">
    <Input
      value={editedProfile.invoiceNumberFormat || ''}
      onChange={e => handleChange('invoiceNumberFormat', e.target.value)}
      placeholder="INV-{YYYY}-{NNNNN}"
      className="font-mono flex-1"
    />
    <div className="flex items-center px-3 rounded border bg-slate-50 text-sm font-mono min-w-[140px]">
      {previewFormat(editedProfile.invoiceNumberFormat || 'INV-{YYYY}-{NNNNN}')}
    </div>
  </div>

  {/* Quick-insert preset buttons */}
  <div className="flex flex-wrap gap-2">
    {[
      { label: 'INV-YYYY-NNNNN', value: 'INV-{YYYY}-{NNNNN}' },
      { label: 'YYYY/NNN',       value: '{YYYY}/{NNN}' },
      { label: 'YYNNNNN',        value: '{YY}{NNNNN}' },
      { label: 'Prefix-only',    value: 'INV-{NNNNN}' },
    ].map(preset => (
      <Button
        key={preset.value}
        variant="outline"
        size="xs"
        onClick={() => handleChange('invoiceNumberFormat', preset.value)}
      >
        {preset.label}
      </Button>
    ))}
  </div>

  {/* Token reference */}
  <p className="text-xs text-slate-500">
    Tokens: {'{YYYY}'} {'{YY}'} {'{MM}'} {'{N}'} {'{NNN}'} {'{NNNNN}'} {'{TENANT}'}
  </p>

  {/* Same block repeated for letterNumberFormat */}
</div>
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 5
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 5 — Admin Role-Based Access Control / RBAC (HIGH)')
divider()

label_value('Status:', 'MISSING — all authenticated admins see all screens; no permission matrix', val_colour=RED)
label_value('Files:', 'src/components/admin/AdminLayout.tsx  |  src/services/adminApi.ts (lines 37–45)  |  src/stores/adminStore.ts')
label_value('Effort:', '3–4 hours')

h2('Current State')
para(
    'AdminLayout.tsx line 17 reads adminUser from Zustand store and renders the full '
    'navigation sidebar regardless of the user\'s role. The Axios request interceptor '
    'in adminApi.ts (lines 37–45) injects the Bearer token correctly but performs zero '
    'permission checks. There is no role field on the admin user object.'
)

h2('Proposed Role Matrix')
add_table(
    ['Screen', 'super_admin', 'billing_admin', 'support_admin'],
    [
        ('Dashboard',        'YES', 'YES', 'YES'),
        ('Packages',         'YES', 'YES', 'NO'),
        ('Users',            'YES', 'NO',  'YES'),
        ('Billing / Revenue','YES', 'YES', 'NO'),
        ('Analytics',        'YES', 'YES', 'NO'),
        ('Tickets',          'YES', 'NO',  'YES'),
        ('Wiki',             'YES', 'NO',  'YES'),
        ('System Settings',  'YES', 'NO',  'NO'),
        ('CMS Pages',        'YES', 'NO',  'NO'),
    ],
    col_widths=[2.2, 1.1, 1.4, 1.4]
)

h2('Proposed Solution')
para('Step 1 — Add role/permissions to adminStore:', bold=True)
code_block(
"""// src/stores/adminStore.ts — extend AdminUser type:
export interface AdminUser {
  id: string;
  email: string;
  name: string;
  role: 'super_admin' | 'billing_admin' | 'support_admin';
  permissions: string[];  // e.g. ['admin.billing.view', 'admin.users.manage']
}
"""
)

para('Step 2 — Add permission constants and guard helper:', bold=True)
code_block(
"""// src/utils/adminPermissions.ts (new file)
export const SCREEN_PERMISSIONS: Record<string, string> = {
  SApackages:  'admin.packages.manage',
  SAASusers:   'admin.users.view',
  SAbilling:   'admin.billing.view',
  SAusage:     'admin.analytics.view',
  SAsettings:  'admin.settings.manage',
  SATickets:   'admin.tickets.manage',
  SAWiki:      'admin.wiki.view',
  SAPages:     'admin.cms.manage',
};

export const canAccess = (screen: string, permissions: string[]): boolean => {
  const required = SCREEN_PERMISSIONS[screen];
  if (!required) return true;  // No restriction = open to all admins
  return permissions.includes(required);
};
"""
)

para('Step 3 — Apply guard in AdminLayout.tsx:', bold=True)
code_block(
"""// AdminLayout.tsx — add guard before rendering screen:
import { canAccess } from '../../utils/adminPermissions';

const { adminUser } = useAdminStore();

// In sidebar navigation — hide/grey-out forbidden items:
const navItems = allNavItems.filter(item =>
  canAccess(item.screen, adminUser?.permissions ?? [])
);

// Before rendering the active screen:
if (!canAccess(currentScreen, adminUser?.permissions ?? [])) {
  return (
    <div className="flex flex-col items-center justify-center h-full gap-3 text-slate-500">
      <ShieldOff className="h-10 w-10 opacity-40" />
      <p className="text-sm">You don't have permission to view this page.</p>
    </div>
  );
}
"""
)

para('Step 4 — Backend: return role + permissions in /admin/me response:', bold=True)
code_block(
"""// api/app/Controllers/AdminAuth.php — me() method:
return $this->respond([
    'id'          => $admin['id'],
    'email'       => $admin['email'],
    'name'        => $admin['name'],
    'role'        => $admin['role'],         // 'super_admin' | 'billing_admin' | ...
    'permissions' => $this->getPermissions($admin['role']),  // array of strings
]);
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 6
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 6 — Two-Factor Authentication / TOTP (MEDIUM)')
divider()

label_value('Status:', 'MISSING — no TOTP library, no DB schema, no auth flow branching', val_colour=RED)
label_value('Files:', 'api/app/Controllers/Auth.php  |  src/components/screens/Login.tsx  |  src/stores/authStore.ts')
label_value('Effort:', '5–6 hours')

h2('Current State')
para(
    'Auth.php login() validates email/password then immediately issues a JWT. Login.tsx '
    'calls onLogin(email, password) and on success stores the token (authStore.ts). '
    'There is no intermediate step for a second factor. The users table has no totp_secret '
    'or totp_enabled column.'
)

h2('Proposed Solution — TOTP Flow')
para('The flow becomes: credentials → check if 2FA enabled → if yes, prompt for code → verify → issue JWT.', bold=True)
doc.add_paragraph()

para('Step 1 — Database migration:', bold=True)
code_block(
"""// new migration: 2026-04-XX-000000_AddTotpToUsers.php
$this->forge->addColumn('users', [
    'totp_secret'  => ['type' => 'VARCHAR', 'constraint' => 64,  'null' => true],
    'totp_enabled' => ['type' => 'TINYINT', 'constraint' => 1,   'default' => 0],
    'totp_backup_codes' => ['type' => 'JSON', 'null' => true],
]);
"""
)

para('Step 2 — Install PHP TOTP library (composer):', bold=True)
code_block('composer require spomky-labs/otphp')

para('Step 3 — Backend: branched login response + TOTP verify endpoint:', bold=True)
code_block(
"""// Auth.php login() — after password validation succeeds:
if ($user['totp_enabled']) {
    // Issue a short-lived interim token (not the real JWT)
    $interimToken = bin2hex(random_bytes(32));
    cache()->save('2fa_interim_' . $interimToken, $user['id'], 300); // 5 min TTL

    return $this->respond([
        'requiresTwoFactor' => true,
        'interimToken'      => $interimToken,
    ]);
}
// else continue to issue full JWT as before

// New endpoint POST /auth/verify-totp:
public function verifyTotp()
{
    $data  = $this->request->getJSON(true);
    $token = $data['interimToken'] ?? '';
    $code  = $data['code'] ?? '';

    $userId = cache()->get('2fa_interim_' . $token);
    if (!$userId) return $this->fail('Session expired or invalid', 401);

    $user   = $this->userModel->find($userId);
    $totp   = \\OTPHP\\TOTP::create($user['totp_secret']);
    if (!$totp->verify($code, null, 1)) {  // 1 = ±30s window
        return $this->fail('Invalid authentication code', 401);
    }

    cache()->delete('2fa_interim_' . $token);
    $jwt = $this->generateToken($user);
    return $this->respond(['token' => $jwt, 'user' => $user]);
}

// New endpoint POST /auth/setup-totp  (to generate QR code for user to scan):
public function setupTotp()
{
    $user = $this->getCurrentUser();
    $totp = \\OTPHP\\TOTP::generate();
    $totp->setLabel($user['email']);
    $totp->setIssuer('BillingTool');

    // Store secret temporarily until user confirms
    cache()->save('2fa_setup_' . $user['id'], $totp->getSecret(), 600);

    return $this->respond([
        'secret'      => $totp->getSecret(),
        'provisionUri'=> $totp->getProvisioningUri(),
        'qrCodeUrl'   => 'https://api.qrserver.com/v1/create-qr-code/?data=' .
                          urlencode($totp->getProvisioningUri()),
    ]);
}
"""
)

para('Step 4 — Frontend: branched login flow in Login.tsx:', bold=True)
code_block(
"""// Login.tsx
const [step, setStep]               = useState<'credentials' | '2fa'>('credentials');
const [interimToken, setInterimToken] = useState('');
const [totpCode, setTotpCode]       = useState('');

const handleLogin = async (e: FormEvent) => {
  e.preventDefault();
  const result = await authService.login(email, password);
  if (result.requiresTwoFactor) {
    setInterimToken(result.interimToken);
    setStep('2fa');
  }
  // else: login complete, store token as usual
};

const handleVerify2FA = async (e: FormEvent) => {
  e.preventDefault();
  await authService.verifyTotp(interimToken, totpCode);
  // authStore will be updated by verifyTotp on success
};

// JSX:
{step === 'credentials' ? (
  <form onSubmit={handleLogin}>
    {/* existing email + password fields */}
  </form>
) : (
  <form onSubmit={handleVerify2FA}>
    <p className="text-sm text-slate-600 mb-3">
      Enter the 6-digit code from your authenticator app.
    </p>
    <Input
      value={totpCode}
      onChange={e => setTotpCode(e.target.value)}
      maxLength={6}
      placeholder="000000"
      className="text-center text-2xl tracking-widest font-mono"
      autoFocus
    />
    <Button type="submit" className="w-full mt-3">Verify</Button>
    <button type="button" className="text-xs text-slate-400 mt-2"
      onClick={() => setStep('credentials')}>
      ← Back
    </button>
  </form>
)}
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 7
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 7 — Invoice Locking After "Sent" Status (MEDIUM)')
divider()

label_value('Status:', 'MISSING — InvoiceEditor has no readOnly mode; backend has no lock check', val_colour=RED)
label_value('Files:', 'src/components/screens/InvoiceEditor.tsx (lines 48–67)  |  api/app/Controllers/InvoiceController.php (update method)')
label_value('Effort:', '2 hours')

h2('Current State')
para(
    'InvoiceEditorProps (line 53) has no readOnly prop. Every field is always editable. '
    'The backend update() method (lines 219–258) performs no lock check before applying changes. '
    'A user can edit an invoice that has already been sent to a customer.'
)

h2('Proposed Solution')
para('Frontend — add isLocked derived state and disabled props:', bold=True)
code_block(
"""// InvoiceEditor.tsx
interface InvoiceEditorProps {
  invoice: Invoice;
  onSave: (invoice: Invoice) => void;
  onBack: () => void;
  onPreview?: () => void;
  mode?: 'invoice' | 'template';
  isTemplate?: boolean;
  // ↓ NEW
  readOnly?: boolean;
}

export function InvoiceEditor({ ..., readOnly = false }: InvoiceEditorProps) {
  const isLocked = readOnly ||
    invoice.status === 'sent'  ||
    invoice.status === 'paid'  ||
    invoice.status === 'cancelled';

  return (
    <>
      {isLocked && (
        <div className="flex items-center gap-2 bg-amber-50 border border-amber-200
                        rounded-md px-4 py-2 mb-4 text-sm text-amber-800">
          <Lock className="h-4 w-4 shrink-0" />
          This {isBusinessLetter ? 'letter' : 'invoice'} is locked and cannot be edited.
        </div>
      )}

      {/* Apply disabled={isLocked} to every <Input>, <Select>, <Textarea> */}
      <Input disabled={isLocked} value={invoice.invoiceNumber} ... />
    </>
  );
}
"""
)

para('Backend — reject edits on locked invoices:', bold=True)
code_block(
"""// InvoiceController.php update() — add at the top of the method:
$existing = $model->find($id);
if (!$existing) return $this->failNotFound();

$lockedStatuses = ['sent', 'paid', 'cancelled'];
if (in_array($existing['status'], $lockedStatuses, true) &&
    // Allow status-only updates (e.g. sent → paid) but block field edits
    count(array_diff_key($dbData, ['status' => true])) > 0)
{
    return $this->fail(
        'Invoice is locked. Only status changes are permitted after it has been ' .
        $existing['status'] . '.',
        403
    );
}
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 8
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 8 — Email Notifications on Invoice Events (MEDIUM)')
divider()

label_value('Status:', 'MISSING — no email service layer, no event hooks', val_colour=RED)
label_value('Files:', 'api/app/Controllers/InvoiceController.php  |  api/app/Config/Email.php (if present)')
label_value('Effort:', '4–5 hours')

h2('Events That Should Trigger Emails')
add_table(
    ['Event', 'Recipient', 'Template'],
    [
        ('Invoice created',   'Seller (internal)',       'invoice_created.html'),
        ('Invoice sent',      'Buyer (customer)',        'invoice_sent.html  + PDF attachment'),
        ('Invoice paid',      'Seller (confirmation)',   'invoice_paid.html'),
        ('Invoice overdue',   'Buyer (reminder)',        'invoice_overdue.html  (cron job)'),
        ('Password reset',    'User',                    'password_reset.html'),
        ('New user signup',   'User + Admin',            'welcome.html'),
    ],
    col_widths=[1.8, 1.8, 2.9]
)

h2('Proposed Solution')
para('Step 1 — Create a reusable EmailService:', bold=True)
code_block(
"""// api/app/Services/EmailService.php
namespace App\\Services;

class EmailService
{
    private static function mailer(): \\CodeIgniter\\Email\\Email
    {
        return service('email');
    }

    public static function sendInvoiceSent(array $invoice, string $buyerEmail, ?string $pdfPath = null): bool
    {
        $email = self::mailer();
        $email->initialize(['mailType' => 'html', 'charset' => 'utf-8']);
        $email->setFrom(env('MAIL_FROM_ADDRESS'), env('MAIL_FROM_NAME', 'BillingTool'));
        $email->setTo($buyerEmail);
        $email->setSubject('Invoice ' . $invoice['invoice_number'] . ' from ' . $invoice['seller_name']);

        $body = view('emails/invoice_sent', ['invoice' => $invoice]);
        $email->setMessage($body);

        if ($pdfPath && file_exists($pdfPath)) {
            $email->attach($pdfPath);
        }

        return $email->send();
    }

    public static function sendInvoicePaid(array $invoice, string $sellerEmail): bool { ... }
    public static function sendOverdueReminder(array $invoice, string $buyerEmail): bool { ... }
}
"""
)

para('Step 2 — Wire into InvoiceController.php update():', bold=True)
code_block(
"""// After status successfully changes to 'sent':
if ($dbData['status'] === 'sent') {
    $buyer = (new BuyerModel())->find($invoice['buyer_id']);
    if ($buyer && !empty($buyer['contact_email'])) {
        \\App\\Services\\EmailService::sendInvoiceSent($invoice, $buyer['contact_email']);
    }
}
"""
)

para('Step 3 — Create HTML email template views/emails/invoice_sent.php:', bold=True)
code_block(
"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"><style>body{font-family:Arial,sans-serif}</style></head>
<body>
  <h2>Invoice <?= $invoice['invoice_number'] ?></h2>
  <p>You have received an invoice from <strong><?= $invoice['seller_name'] ?></strong>.</p>
  <table>
    <tr><td>Amount:</td><td><?= $invoice['currency'] ?> <?= number_format($invoice['grand_total'],2) ?></td></tr>
    <tr><td>Due:</td><td><?= $invoice['due_date'] ?></td></tr>
  </table>
  <p>Please find the PDF attached.</p>
</body>
</html>
"""
)

para('Step 4 — Configure SMTP in api/.env:', bold=True)
code_block(
"""email.protocol  = smtp
email.SMTPHost  = smtp.mailgun.org
email.SMTPUser  = postmaster@mg.yourdomain.com
email.SMTPPass  = your-api-key
email.SMTPPort  = 587
email.SMTPCrypto = tls
email.fromEmail = noreply@billingtool.com
email.fromName  = BillingTool
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 9
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 9 — Batch PDF Export (MEDIUM)')
divider()

label_value('Status:', 'PARTIAL — multi-select UI and bulk export button exist; exportInvoicesBulk() is not implemented', val_colour=AMBER)
label_value('Files:', 'src/components/screens/InvoiceList.tsx (lines 169–188, 553–586)  |  src/utils/invoice-pdf.ts')
label_value('Effort:', '3–4 hours')

h2('Current State')
para(
    'InvoiceList.tsx has selectedInvoices state (Set<string>) and checkboxes in each row '
    '(lines 596–600, 987–990). A bulk actions toolbar (lines 553–586) contains an "Export" '
    'button that calls handleBulkExport() (lines 169–188), which in turn calls '
    'exportInvoicesBulk() — but that utility function is either not implemented or not '
    'returning a ZIP download.'
)

h2('Proposed Solution')
para('Step 1 — Install JSZip:', bold=True)
code_block('npm install jszip')

para('Step 2 — Implement exportInvoicesBulk in src/utils/invoice-export.ts:', bold=True)
code_block(
"""// src/utils/invoice-export.ts
import JSZip from 'jszip';
import { Invoice } from '../types/invoice';
import { generateInvoicePDF } from './invoice-pdf';
import { generateUBLXML } from './invoice-ubl';

export type BulkExportFormat = 'pdf' | 'ubl-xml' | 'json' | 'csv';

export async function exportInvoicesBulk(
    invoices: Invoice[],
    format: BulkExportFormat,
    onProgress?: (pct: number) => void
): Promise<void> {
    const zip = new JSZip();
    const folder = zip.folder('invoices')!;

    for (let i = 0; i < invoices.length; i++) {
        const inv = invoices[i];
        onProgress?.(Math.round((i / invoices.length) * 90));

        const safeNum = inv.invoiceNumber.replace(/[^a-zA-Z0-9_-]/g, '_');

        if (format === 'pdf') {
            const pdfBlob = await generateInvoicePDF(inv, { returnBlob: true });
            folder.file(`${safeNum}.pdf`, pdfBlob as Blob);
        } else if (format === 'ubl-xml') {
            const xml = generateUBLXML(inv);
            folder.file(`${safeNum}.xml`, xml);
        } else if (format === 'json') {
            folder.file(`${safeNum}.json`, JSON.stringify(inv, null, 2));
        } else if (format === 'csv') {
            const row = [inv.invoiceNumber, inv.status, inv.grandTotal, inv.currency].join(',');
            folder.file(`${safeNum}.csv`, row);
        }
    }

    onProgress?.(95);
    const blob = await zip.generateAsync({ type: 'blob' });
    onProgress?.(100);

    const url = URL.createObjectURL(blob);
    const a   = document.createElement('a');
    a.href    = url;
    a.download = `export_${new Date().toISOString().split('T')[0]}.zip`;
    a.click();
    URL.revokeObjectURL(url);
}
"""
)

para('Step 3 — Wire progress bar in InvoiceList handleBulkExport:', bold=True)
code_block(
"""// InvoiceList.tsx — handleBulkExport:
const [exportProgress, setExportProgress] = useState(0);

const handleBulkExport = async () => {
  const toExport = invoices.filter(inv => selectedInvoices.has(inv.id));
  toast.info(`Exporting ${toExport.length} invoices…`);
  await exportInvoicesBulk(toExport, exportFormat, (pct) => setExportProgress(pct));
  toast.success('Download ready');
  setSelectedInvoices(new Set());
};
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 10
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 10 — Business Letter Body Rich-Text Editor (MEDIUM)')
divider()

label_value('Status:', 'PARTIAL — isBusinessLetter flag exists; invoice.body field is typed; no rich-text UI is rendered for letter mode', val_colour=AMBER)
label_value('Files:', 'src/components/screens/InvoiceEditor.tsx (line 62)  |  src/components/ui/RichTextEditor.tsx  |  src/types/invoice.ts (line 88)')
label_value('Effort:', '1.5–2 hours')

h2('Current State')
para(
    'InvoiceEditor.tsx line 62 sets isBusinessLetter = invoice.templateType === "business_letter". '
    'The Invoice type (invoice.ts:88) has a body?: string field. However, no UI conditionally '
    'renders a rich-text editor for the letter body when isBusinessLetter is true. The only '
    'letter-specific field visible is the Subject (a plain <Input> mapping to invoice.note).'
)
para(
    'RichTextEditor.tsx is a fully featured TipTap editor (bold, italic, underline, bullet lists, '
    'headings) and is already used in Settings.tsx for headerText / footerText fields.'
)

h2('Proposed Solution')
code_block(
"""// InvoiceEditor.tsx — add after the subject/note section:

// 1. Import the editor at the top:
import { RichTextEditor } from '../ui/RichTextEditor';

// 2. Add state (or derive from invoice):
const [letterBody, setLetterBody] = useState(invoice.body || '');

// 3. Sync state back to invoice:
const handleLetterBodyChange = (html: string) => {
  setLetterBody(html);
  handleUpdateInvoice({ body: html });
};

// 4. Render below the recipient section, only for business letters:
{isBusinessLetter && (
  <div className="rounded-xl border bg-card p-5 space-y-3">
    <h3 className="text-sm font-semibold text-slate-700">
      {t('editor.letterBody') || 'Letter Body'}
    </h3>
    <div className="min-h-[260px] rounded-md border bg-white">
      <RichTextEditor
        value={letterBody}
        onChange={handleLetterBodyChange}
        placeholder={
          t('editor.letterBodyPlaceholder') ||
          'Dear Sir/Madam,\\n\\nWrite your letter content here...'
        }
      />
    </div>
  </div>
)}

// 5. Make sure PDF picks up invoice.body — verify invoice-pdf.ts line ~406:
//    if (el.type === 'description') { doc.text(invoice.body || '', ...) }
//    Add HTML-to-plain-text stripping if needed:
import { htmlToText } from 'html-to-text';  // npm install html-to-text
const bodyText = htmlToText(invoice.body || '');
"""
)

para('Also add the translation key to all 4 language files:', bold=True)
code_block(
"""// en.ts, de.ts, ar.ts, pl.ts — add inside the editor section:
letterBody: 'Letter Body',                    // en
letterBody: 'Brieftext',                      // de
letterBody: 'نص الخطاب',                      // ar
letterBody: 'Treść listu',                    // pl
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 11
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 11 — customerApi.ts Missing Request Interceptor (LOW)')
divider()

label_value('Status:', 'PARTIAL — every method manually passes Authorization headers instead of using an interceptor', val_colour=AMBER)
label_value('Files:', 'src/services/customerApi.ts (lines 1–116)  |  src/services/api.ts (lines 16–29)')
label_value('Effort:', '1–1.5 hours')

h2('Current State')
para(
    'api.ts (lines 16–29) has a request interceptor that reads the token from Zustand '
    'authStore and injects it automatically. customerApi.ts has only a response interceptor '
    '(401 redirect) but no request interceptor — so every single method call includes:'
)
code_block(
"""// customerApi.ts (lines 46–54) — repeated in every method:
const response = await customerApi.get('/dashboard', {
  headers: {
    'Authorization': `Bearer ${token}`,
    'X-Authorization': `Bearer ${token}`,
  }
});
"""
)

h2('Proposed Solution')
code_block(
"""// customerApi.ts — add request interceptor after the axios.create() call:

import { useAuthStore } from '../stores/authStore';

customerApi.interceptors.request.use((config) => {
  const token = useAuthStore.getState().token;
  if (token) {
    config.headers['Authorization']   = `Bearer ${token}`;
    config.headers['X-Authorization'] = `Bearer ${token}`;
  }
  return config;
});

// Then simplify every method — no more manual headers:
export const customerService = {
  getDashboard: async () => {
    const response = await customerApi.get<{ success: boolean; data: DashboardData }>('/dashboard');
    return response.data.data;
  },

  getInvoices: async (page = 1, perPage = 10) => {
    const response = await customerApi.get('/invoices', { params: { page, per_page: perPage } });
    return response.data;
  },

  // Repeat for all other methods — remove the { headers: { ... } } option
};
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 12
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 12 — Wiki Admin Edit UI (LOW)')
divider()

label_value('Status:', 'MISSING — SAWiki.tsx is a read-only viewer; no edit mode, no save button, no updateContent API method', val_colour=RED)
label_value('Files:', 'src/components/screens/Admin/SAWiki.tsx (lines 116–558)  |  src/services/adminApi.ts (wiki section)')
label_value('Effort:', '3–4 hours')

h2('Current State')
para(
    'SAWiki.tsx renders a file-tree sidebar and Markdown viewer. Buttons are: Export PDF. '
    'There is no Edit button, no textarea/editor, and no save path. adminWikiService in '
    'adminApi.ts exposes getTree() and getContent() but no updateContent().'
)

h2('Proposed Solution')
para('Step 1 — Add updateContent to adminApi.ts:', bold=True)
code_block(
"""// src/services/adminApi.ts — extend adminWikiService:
export const adminWikiService = {
  getTree:    (language: string) =>
    adminApi.get('/wiki', { params: { language } }).then(r => r.data),
  getContent: (path: string, language: string) =>
    adminApi.get('/wiki/content', { params: { path, language } }).then(r => r.data),
  // NEW:
  updateContent: (path: string, content: string, language: string) =>
    adminApi.put('/wiki/content', { path, content, language }).then(r => r.data),
};
"""
)

para('Step 2 — Add backend endpoint in AdminWiki.php:', bold=True)
code_block(
"""// api/app/Controllers/Admin/AdminWiki.php
public function updateContent()
{
    $data     = $this->request->getJSON(true);
    $path     = $data['path']     ?? '';
    $content  = $data['content']  ?? '';
    $language = $data['language'] ?? 'en';

    if (empty($path)) return $this->fail('Path is required');

    // Sanitize path to prevent directory traversal
    $safePath = realpath(WRITEPATH . 'wiki/' . $language . '/' . ltrim($path, '/'));
    if (!$safePath || !str_starts_with($safePath, realpath(WRITEPATH . 'wiki/'))) {
        return $this->fail('Invalid path', 400);
    }

    file_put_contents($safePath, $content);
    return $this->respondUpdated(['message' => 'Saved']);
}
"""
)

para('Step 3 — Add route in Routes.php:', bold=True)
code_block("""$routes->put('admin/wiki/content', 'Admin\\AdminWiki::updateContent');""")

para('Step 4 — Add edit mode state and UI to SAWiki.tsx:', bold=True)
code_block(
"""// SAWiki.tsx — add state:
const [isEditing,   setIsEditing]   = useState(false);
const [editContent, setEditContent] = useState('');

// Header buttons:
{selectedPath && (
  <>
    <Button variant="outline" size="sm" onClick={() => {
      setEditContent(content); setIsEditing(true);
    }}>
      <Pencil className="h-3.5 w-3.5 mr-1" /> Edit
    </Button>
    {isEditing && (
      <Button size="sm" onClick={async () => {
        await adminWikiService.updateContent(selectedPath, editContent, language);
        toast.success('Page saved');
        setIsEditing(false);
        setContent(editContent);
      }}>
        Save
      </Button>
    )}
  </>
)}

// Viewer area:
{isEditing ? (
  <textarea
    value={editContent}
    onChange={e => setEditContent(e.target.value)}
    className="w-full h-full p-4 font-mono text-sm resize-none border-none outline-none"
  />
) : (
  <ReactMarkdown>{content}</ReactMarkdown>
)}
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 13
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 13 — PDF/A-3 with Embedded UBL XML (LOW)')
divider()

label_value('Status:', 'MISSING — embedPdfInUbl property exists in ExportOptions type but is never used; jsPDF cannot produce PDF/A-3', val_colour=RED)
label_value('Files:', 'src/utils/invoice-pdf.ts (lines 1–30)  |  src/types/invoice.ts (lines 171–177)')
label_value('Effort:', '6–8 hours (server-side implementation required)')

h2('Why It Cannot Be Done Client-Side')
para(
    'PDF/A-3 (ISO 19005-3) requires embedding file attachments inside the PDF binary and '
    'setting conformance metadata (XMP). The jsPDF library used in invoice-pdf.ts does not '
    'support PDF/A-3 conformance or embedded file attachments. This must be implemented '
    'server-side using a PDF library like TCPDF or mPDF (PHP) or PDFBox (Java).'
)

h2('Proposed Solution — Backend Endpoint')
para('Step 1 — Add TCPDF via Composer:', bold=True)
code_block('composer require tecnickcom/tcpdf')

para('Step 2 — Create a new controller/method:', bold=True)
code_block(
"""// api/app/Controllers/InvoiceController.php — add method:
public function exportPDFA3($id = null)
{
    $model   = new InvoiceModel();
    $invoice = $model->find($id);
    if (!$invoice) return $this->failNotFound();

    // 1. Generate UBL XML
    $ublXml = $this->generateUBL($invoice);

    // 2. Render HTML invoice template
    $html = view('pdf/invoice', ['invoice' => $invoice]);

    // 3. Create PDF/A-3 via TCPDF
    $pdf = new \\TCPDF('P', 'mm', 'A4', true, 'UTF-8');
    $pdf->SetCreator('BillingTool');
    $pdf->SetTitle('Invoice ' . $invoice['invoice_number']);
    $pdf->setPDFVersion('1.7');
    $pdf->AddPage();
    $pdf->writeHTML($html);

    // 4. Attach UBL XML (makes it PDF/A-3)
    $pdf->addFileAttachment(
        $invoice['invoice_number'] . '.xml',
        $ublXml,
        'UBL 2.1 Invoice',
        'application/xml'
    );

    $filename = $invoice['invoice_number'] . '_PDF-A3.pdf';
    return $this->response
        ->setHeader('Content-Type', 'application/pdf')
        ->setHeader('Content-Disposition', "attachment; filename=\\"{$filename}\\"")
        ->setBody($pdf->Output('', 'S'));
}
"""
)

para('Step 3 — Add export button in InvoicePreview.tsx / InvoiceList.tsx:', bold=True)
code_block(
"""<Button variant="outline" size="sm" onClick={() => {
  window.location.href = `/api/invoices/${invoice.id}/export-pdfa3`;
}}>
  PDF/A-3 + UBL
</Button>
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 14
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 14 — Invoice Clone/Copy Persisted to Database (LOW)')
divider()

label_value('Status:', 'PARTIAL — handleDuplicate() exists in InvoiceList.tsx (lines 230–243) but only mutates local state; not saved to backend', val_colour=AMBER)
label_value('Files:', 'src/components/screens/InvoiceList.tsx (lines 230–243)')
label_value('Effort:', '1–1.5 hours')

h2('Current State')
code_block(
"""// InvoiceList.tsx — current handleDuplicate (lines 230–243):
const handleDuplicate = (invoice: Invoice) => {
  const newInvoice: Invoice = {
    ...invoice,
    id: `new_${Math.random().toString(36).substring(7)}`,  // Fake ID
    invoiceNumber: `${invoice.invoiceNumber}-COPY`,
    status: 'draft',
    createdAt: new Date().toISOString(),
    updatedAt: new Date().toISOString(),
  };
  setInvoices([newInvoice, ...invoices]);  // Local state only — lost on refresh!
  toast.success(`Invoice ${invoice.invoiceNumber} duplicated`);
};
"""
)

h2('Proposed Solution')
code_block(
"""// InvoiceList.tsx — replace handleDuplicate:
const handleDuplicate = async (invoice: Invoice) => {
  setIsLoading(true);
  try {
    const clone: Partial<Invoice> = {
      ...invoice,
      id:            undefined,   // Let backend auto-assign
      invoiceNumber: `${invoice.invoiceNumber}-COPY`,
      status:        'draft',
      issueDate:     new Date().toISOString().split('T')[0],
      dueDate:       undefined,
      createdAt:     undefined,
      updatedAt:     undefined,
    };

    const saved = await invoiceService.create(clone as Invoice);
    setInvoices(prev => [saved, ...prev]);

    toast.success(t('invoiceList.duplicated') || 'Invoice duplicated', {
      description: `Saved as ${clone.invoiceNumber}`,
      action: {
        label: 'Edit',
        onClick: () => onEditInvoice(saved),
      },
    });
  } catch {
    toast.error(t('invoiceList.duplicateFailed') || 'Could not duplicate invoice');
  } finally {
    setIsLoading(false);
  }
};
"""
)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# GAP 15
# ─────────────────────────────────────────────────────────────────────────────
h1('Gap 15 — AuditLog Field-Level Change Tracking (LOW)')
divider()

label_value('Status:', 'PARTIAL — logAction() trait is called on status changes but records no field-level diff or user context', val_colour=AMBER)
label_value('Files:', 'api/app/Controllers/InvoiceController.php  |  api/app/Models/AuditLogModel.php')
label_value('Effort:', '2–3 hours')

h2('Current State')
para(
    'AuditLogModel.php allowedFields: tenant_id, timestamp, action, invoice_number, user, '
    'details, signed. logAction() is called after status changes with a generic string '
    '("Invoice sent to buyer"). No field-level diff is captured, and the user field is '
    'not populated from the authenticated JWT.'
)

h2('Proposed Solution')
para('Step 1 — Add a JSON details column for field diffs (if not already JSON):', bold=True)
code_block(
"""// Migration: 2026-04-XX-AddAuditDetails.php
$this->forge->modifyColumn('audit_logs', [
    'details' => ['type' => 'JSON', 'null' => true],
    'changed_fields' => ['type' => 'JSON', 'null' => true],  // NEW
    'user_id'        => ['type' => 'INT', 'null' => true],   // NEW
]);
"""
)

para('Step 2 — Enhance logAction in AuditTrait.php:', bold=True)
code_block(
"""// api/app/Traits/AuditTrait.php
protected function logAction(
    string $action,
    string $invoiceNumber,
    string $details = '',
    bool $signed = false,
    array $changedFields = []
): void {
    $userId = $this->getAuthenticatedUserId();

    (new \\App\\Models\\AuditLogModel())->insert([
        'tenant_id'      => $this->getAuthenticatedTenantId(),
        'timestamp'      => date('Y-m-d H:i:s'),
        'action'         => $action,
        'invoice_number' => $invoiceNumber,
        'user'           => $userId,
        'user_id'        => $userId,
        'details'        => $details,
        'changed_fields' => empty($changedFields) ? null : json_encode($changedFields),
        'signed'         => $signed,
    ]);
}

private function getAuthenticatedUserId(): ?int
{
    $header = service('request')->getHeaderLine('Authorization');
    if (preg_match('/Bearer\\s(\\S+)/', $header, $m)) {
        try {
            $key     = env('JWT_SECRET');
            $decoded = \\Firebase\\JWT\\JWT::decode($m[1], new \\Firebase\\JWT\\Key($key, 'HS256'));
            return $decoded->uid ?? null;
        } catch (\\Exception $e) {}
    }
    return null;
}
"""
)

para('Step 3 — Call with field diff in InvoiceController update():', bold=True)
code_block(
"""// InvoiceController.php update() — capture diff before update:
$before = $model->find($id);

if ($model->update($id, $dbData)) {
    $after = $model->find($id);

    // Compute diff for auditable fields only
    $auditFields   = ['status','buyer_id','grand_total','due_date','currency'];
    $changedFields = [];
    foreach ($auditFields as $field) {
        if ($before[$field] !== $after[$field]) {
            $changedFields[$field] = ['from' => $before[$field], 'to' => $after[$field]];
        }
    }

    $this->logAction('updated', $after['invoice_number'], 'Invoice updated', false, $changedFields);
}
"""
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDIX — EFFORT SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1('Appendix — Consolidated Effort Summary')
divider()

effort_rows = [
    ('1',  'Template Designer drag-drop',          'HIGH',   '3–4 h',   'Frontend only  ·  TemplateDesignLayout.tsx'),
    ('2',  'Status transition enforcement',        'HIGH',   '2 h',     'Backend only  ·  InvoiceController.php'),
    ('3',  'Logo upload in Settings',              'HIGH',   '1–1.5 h', 'Frontend only  ·  Settings.tsx'),
    ('4',  'Invoice number format builder',        'HIGH',   '2–2.5 h', 'Frontend only  ·  Settings.tsx'),
    ('5',  'Admin RBAC',                           'HIGH',   '3–4 h',   'Both  ·  AdminLayout + Auth.php'),
    ('6',  '2FA / TOTP',                           'MEDIUM', '5–6 h',   'Both  ·  Auth.php + Login.tsx + DB'),
    ('7',  'Invoice locking after "sent"',         'MEDIUM', '2 h',     'Both  ·  InvoiceEditor + Controller'),
    ('8',  'Email notifications',                  'MEDIUM', '4–5 h',   'Backend only  ·  EmailService + SMTP'),
    ('9',  'Batch PDF export',                     'MEDIUM', '3–4 h',   'Frontend only  ·  invoice-export.ts'),
    ('10', 'Letter body rich-text editor',         'MEDIUM', '1.5–2 h', 'Frontend only  ·  InvoiceEditor.tsx'),
    ('11', 'customerApi interceptor',              'LOW',    '1–1.5 h', 'Frontend only  ·  customerApi.ts'),
    ('12', 'Wiki admin edit UI',                   'LOW',    '3–4 h',   'Both  ·  SAWiki.tsx + AdminWiki.php'),
    ('13', 'PDF/A-3 with embedded UBL',            'LOW',    '6–8 h',   'Backend only  ·  TCPDF + InvoiceCtrl'),
    ('14', 'Invoice clone persisted to DB',        'LOW',    '1–1.5 h', 'Frontend only  ·  InvoiceList.tsx'),
    ('15', 'AuditLog field-level tracking',        'LOW',    '2–3 h',   'Backend only  ·  AuditTrait.php'),
]

add_table(
    ['#', 'Gap', 'Priority', 'Effort', 'Scope'],
    effort_rows,
    col_widths=[0.25, 2.35, 0.80, 0.85, 2.25]
)

para('')
para('Total minimum estimate:  39.5 hours', bold=True, colour=PURPLE)
para('Total maximum estimate:  60.0 hours', bold=True, colour=PURPLE)
para('')
para(
    'Sprint recommendation: tackle Gaps 1–5 (HIGH priority) first — together they represent '
    '~12 hours and unlock the most visible user-facing improvements. '
    'Gaps 6 and 8 (2FA and Email) are medium priority but require the most cross-stack work.',
    colour=GREY, size=10
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/sivaji/Downloads/BillingTool/BillingTool_Detailed_Gap_Analysis.docx'
doc.save(out)
print(f'Saved: {out}')
