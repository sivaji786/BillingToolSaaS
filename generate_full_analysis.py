from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────
s = doc.sections[0]
s.page_width = Inches(8.5); s.page_height = Inches(11)
s.left_margin = s.right_margin = Inches(0.85)
s.top_margin = s.bottom_margin = Inches(0.85)

# ── Palette ────────────────────────────────────────────────────────
PURPLE = RGBColor(0x6B,0x21,0xA8)
DARK   = RGBColor(0x1F,0x2A,0x3C)
GREY   = RGBColor(0x6B,0x72,0x80)
GREEN  = RGBColor(0x05,0x60,0x27)
AMBER  = RGBColor(0x92,0x40,0x0E)
RED    = RGBColor(0x99,0x17,0x17)
BLUE   = RGBColor(0x1E,0x40,0xAF)
WHITE  = RGBColor(0xFF,0xFF,0xFF)

# ── Helpers ────────────────────────────────────────────────────────
def shd(cell, hex6):
    e = OxmlElement('w:shd')
    e.set(qn('w:val'),'clear'); e.set(qn('w:color'),'auto'); e.set(qn('w:fill'),hex6)
    cell._tc.get_or_add_tcPr().append(e)

def divider(color='6B21A8'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1');    b.set(qn('w:color'),color)
    pBdr.append(b); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(5)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = PURPLE; r.font.bold = True; r.font.size = Pt(16)

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = PURPLE; r.font.bold = True; r.font.size = Pt(13)

def h3(text, colour=DARK):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = colour; r.font.bold = True; r.font.size = Pt(11)

def para(text, bold=False, colour=DARK, size=10.5, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Inches(indent*0.22)
    r = p.add_run(text)
    r.font.bold=bold; r.font.color.rgb=colour; r.font.size=Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, colour=DARK, size=10, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent*0.22)
    r = p.add_run(text)
    r.font.color.rgb=colour; r.font.size=Pt(size)
    p.paragraph_format.space_after = Pt(3)

def code(text):
    t = doc.add_table(rows=1,cols=1)
    t.style='Table Grid'
    c = t.rows[0].cells[0]
    shd(c,'F1F5F9')
    p = c.paragraphs[0]
    r = p.add_run(text)
    r.font.name='Courier New'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(0x1E,0x29,0x3B)
    doc.add_paragraph()

def status_colour(txt):
    t=str(txt)
    if any(k in t for k in ('DONE','COMPLETE','Working')): return GREEN
    if any(k in t for k in ('PARTIAL','Partial','STUB','partial')): return AMBER
    if any(k in t for k in ('MISSING','Missing','NOT IMPL','ABANDONED','BROKEN','CRITICAL')): return RED
    if any(k in t for k in ('HIGH',)): return RED
    if any(k in t for k in ('MEDIUM',)): return AMBER
    if any(k in t for k in ('LOW',)): return GREEN
    return DARK

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; shd(c,'6B21A8')
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        ru=p.add_run(h); ru.font.bold=True; ru.font.color.rgb=WHITE; ru.font.size=Pt(9)
    for ri,row in enumerate(rows):
        nr=t.add_row(); bg='F5F3FF' if ri%2 else 'FFFFFF'
        for ci,val in enumerate(row):
            c=nr.cells[ci]; shd(c,bg)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            p=c.paragraphs[0]
            ru=p.add_run(str(val))
            ru.font.size=Pt(9); ru.font.color.rgb=status_colour(val)
    if widths:
        for ci,w in enumerate(widths):
            for row in t.rows: row.cells[ci].width=Inches(w)
    doc.add_paragraph()
    return t

def lv(label, val, lc=PURPLE, vc=DARK):
    p=doc.add_paragraph()
    r1=p.add_run(label+'  '); r1.font.bold=True; r1.font.color.rgb=lc; r1.font.size=Pt(10)
    r2=p.add_run(val); r2.font.color.rgb=vc; r2.font.size=Pt(10)
    p.paragraph_format.space_after=Pt(3)

# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(72)
r=p.add_run('BillingTool'); r.font.bold=True; r.font.size=Pt(36); r.font.color.rgb=PURPLE

p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=p2.add_run('Complete Application Analysis — All Modules')
r2.font.size=Pt(17); r2.font.color.rgb=DARK

p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r3=p3.add_run('Implementation Status · Gaps · Proposed Solutions · Effort Estimates')
r3.font.size=Pt(11); r3.font.color.rgb=GREY

p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
r4=p4.add_run('Generated: 2026-04-23'); r4.font.size=Pt(10); r4.font.color.rgb=GREY
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1 — APPLICATION OVERVIEW
# ══════════════════════════════════════════════════════════════════
h1('1. Application Overview')
divider()
para('BillingTool is a multi-tenant SaaS invoicing and business communication platform. '
     'The frontend is a React 18 + TypeScript SPA (Vite) with hash-based routing. '
     'The backend is CodeIgniter 4 (PHP 8.1) with a MySQL database. '
     'The system supports invoice creation, business letters, file workspace, buyer management, '
     'multi-language UI (EN/DE/AR/PL), admin portal, customer billing/subscriptions, and AI-powered features.')

h2('1.1  Overall Completion by Module')
tbl(
    ['Module','Screens','Backend Controllers','Status','Score'],
    [
        ('Invoice & Letters',    '5',  '2',  'DONE',    '9/10'),
        ('Template System',      '3',  '1',  'PARTIAL', '7/10'),
        ('Buyer Management',     '1',  '1',  'DONE',    '9/10'),
        ('Company Settings',     '1',  '1',  'PARTIAL', '6/10'),
        ('Dashboard & Analytics','1',  '1',  'DONE',    '8/10'),
        ('Workspace / Files',    '2',  '1',  'PARTIAL', '7/10'),
        ('Customer Billing',     '2',  '1',  'PARTIAL', '6/10'),
        ('Authentication',       '3',  '1',  'PARTIAL', '7/10'),
        ('Quick Access Portal',  '1',  '1',  'PARTIAL', '5/10'),
        ('Admin — Dashboard',    '1',  '1',  'PARTIAL', '6/10'),
        ('Admin — Packages',     '3',  '2',  'PARTIAL', '6/10'),
        ('Admin — Users',        '2',  '1',  'PARTIAL', '5/10'),
        ('Admin — Billing',      '2',  '1',  'PARTIAL', '4/10'),
        ('Admin — Analytics',    '1',  '1',  'PARTIAL', '5/10'),
        ('Admin — Settings',     '1',  '1',  'PARTIAL', '6/10'),
        ('Admin — Tickets',      '2',  '1',  'PARTIAL', '6/10'),
        ('Admin — Wiki',         '1',  '1',  'PARTIAL', '5/10'),
        ('Admin — CMS Pages',    '1',  '1',  'PARTIAL', '6/10'),
        ('Multi-language',       '—',  '—',  'DONE',    '9/10'),
        ('Legal / CMS Pages',    '4',  '1',  'DONE',    '8/10'),
        ('Activity / Audit Log', '1',  '1',  'DONE',    '8/10'),
        ('AI Invoice Assistant', '1',  '1',  'PARTIAL', '6/10'),
        ('RBAC / Permissions',   '4',  '3',  'PARTIAL', '5/10'),
        ('OVERALL',              '46+','33+','PARTIAL', '6.8/10'),
    ],
    widths=[2.1,0.8,1.6,0.9,0.6]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 2 — MODULES FULLY COMPLETE
# ══════════════════════════════════════════════════════════════════
h1('2. Fully Implemented Modules')
divider()
para('The following modules are production-ready. No critical gaps were found.')

h2('2.1  Invoice Module')
lv('Files:','src/components/screens/InvoiceEditor.tsx · InvoiceList.tsx · InvoicePreview.tsx · api/app/Controllers/InvoiceController.php')
lv('Status:','DONE — 9/10',vc=GREEN)
for item in [
    'Full CRUD: create, read, update, delete with all validation',
    'EN 16931 compliance validation (src/utils/invoice-validation.ts)',
    'Tax calculations: net, VAT, grand total (src/utils/invoice-calculations.ts)',
    'Status workflow: draft → validated → sent → paid / cancelled',
    'Client-side PDF export with layout from template (src/utils/invoice-pdf.ts)',
    'UBL 2.1 XML, JSON, CSV export',
    'InvoiceList: search, multi-column sort, status filter, date range filter, pagination',
    'Bulk selection with checkboxes; bulk status change and bulk delete',
    'Import: CSV, JSON, UBL XML (src/utils/invoice-import.ts)',
    'Template application: creates a new invoice pre-filled from a saved template',
    'Keyboard shortcuts in editor (Ctrl+S save, Ctrl+P preview)',
    'React Query caching with automatic refetch on mutation',
]:
    bullet(item, colour=GREEN)

para('Minor known gaps (not blocking):')
for item in [
    'No status transition enforcement in backend — any → any is allowed',
    'No invoice locking after "sent" (fields remain editable)',
    'invoice.attachments array typed but no file upload UI',
    'invoice.signed field typed but no signature capture UI',
    'handleDuplicate() creates a local copy but does not persist to backend (API call missing)',
]:
    bullet(item, colour=AMBER)

h2('2.2  Business Letter Module')
lv('Status:','PARTIAL — 6/10',vc=AMBER)
para('Implemented alongside the Invoice module using a templateType flag. Core flow works.')
for item in [
    'templateType: "business_letter" on Invoice and InvoiceTemplate types',
    'DEFAULT_LETTER_LAYOUT in invoice-templates-defaults.ts',
    'handleNewBusinessLetter() handler in App.tsx',
    'InvoiceEditor hides currency, tax, due-date fields when isBusinessLetter',
    'Subject field replaces Due Date (mapped to invoice.note)',
    'PDF renders description element as letter body',
    'List view with tab filter between invoices and letters',
]:
    bullet(item, colour=GREEN)
para('Gaps:')
for item in [
    'No rich-text editor for letter body — only a plain <Input> for the subject/note',
    'invoice.body field typed but never exposed in the editor UI',
    'No salutation, closing, or reference-number fields',
    'Auto-numbering not wired to letterNumberFormat from company profile',
    'Letter preview page still shows "Invoice Preview" heading in some places',
]:
    bullet(item, colour=AMBER)

h2('2.3  Buyer Management')
lv('Files:','src/components/screens/Buyers.tsx · api/app/Controllers/BuyerController.php')
lv('Status:','DONE — 9/10',vc=GREEN)
for item in [
    'Full CRUD with confirmation dialogs on delete',
    'Search, multi-column sort, pagination',
    'Address stored as JSON (migration 2026-02-16)',
    'Buyer autocomplete in InvoiceEditor (BuyerAutocomplete.tsx)',
    'VAT ID, contact email, phone all supported',
]:
    bullet(item, colour=GREEN)

h2('2.4  Dashboard')
lv('Files:','src/components/screens/Dashboard.tsx · (stats from InvoiceController/index)')
lv('Status:','DONE — 8/10',vc=GREEN)
for item in [
    'Revenue cards: total, paid, pending, draft',
    'Status distribution pie chart (recharts)',
    'Monthly revenue trend line chart (last 6 months)',
    'Recent invoices list with status badges',
    'Quick actions: New Invoice, Import, Validate Batch',
    'Import dialog with file drag-and-drop',
    'Batch validation: validates all draft invoices against EN 16931 rules',
]:
    bullet(item, colour=GREEN)
para('Minor gaps: no custom date range picker, no buyer-level analytics.')

h2('2.5  Multi-language (EN / DE / AR / PL)')
lv('Files:','src/translations/{en,de,ar,pl}.ts · src/utils/i18n.ts · src/components/LanguageSwitcher.tsx')
lv('Status:','DONE — 9/10',vc=GREEN)
for item in [
    'Four complete language files: EN (1326 lines), DE (1326), AR (1326), PL (1333)',
    'RTL layout support for Arabic via isRtl hook',
    'Automatic fallback to English for any missing key',
    'LanguageSwitcher dropdown with flag + code label',
    'Language persisted via LanguageContext',
]:
    bullet(item, colour=GREEN)

h2('2.6  Legal / CMS Pages')
lv('Files:','src/components/screens/{PrivacyPolicy,TermsAndConditions,CookiePolicy,Impressum}.tsx · api/app/Controllers/CmsController.php')
lv('Status:','DONE — 8/10',vc=GREEN)
for item in [
    'Four static pages seeded via migration (2026-04-18)',
    'Public API endpoint: GET /api/public/cms/:slug',
    'Admin endpoint: GET /admin/cms + PUT /admin/cms (for editing)',
    'All pages linked from LandingPage navigation',
    'No authentication required to view',
]:
    bullet(item, colour=GREEN)

h2('2.7  Activity / Audit Log')
lv('Files:','src/components/screens/ActivityLog.tsx · api/app/Controllers/AuditLogController.php · api/app/Models/AuditLogModel.php')
lv('Status:','DONE — 8/10',vc=GREEN)
for item in [
    'AuditLogModel with tenant_id, timestamp, action, invoice_number, user, details, signed',
    'logAction() trait used in InvoiceController on create/update/send',
    'ActivityLog screen: search, filter by action, date range, pagination',
    'AuditLogController: paginated index and show endpoints',
]:
    bullet(item, colour=GREEN)
para('Gap: logAction() does not capture field-level diffs or the authenticated user ID (user field is empty).',colour=AMBER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 3 — PARTIALLY IMPLEMENTED MODULES
# ══════════════════════════════════════════════════════════════════
h1('3. Partially Implemented Modules — Detail & Proposed Solutions')
divider()

# ── 3.1 Template System ──────────────────────────────────────────
h2('3.1  Template System')
lv('Files:','src/components/screens/TemplateLibrary.tsx · src/components/invoice/TemplateEditor.tsx · src/components/invoice/TemplateDesignLayout.tsx · src/pages/DesignLayoutPage.tsx')
lv('Status:','PARTIAL — 7/10',vc=AMBER)

h3('What Works')
for item in [
    'TemplateLibrary: lists all templates, tab filter invoice vs business_letter, create/edit/delete',
    'TemplateEditor: form for name, description, logo upload, colors, header/footer text, tax defaults',
    'Template type selector (invoice / business_letter) with default layout assignment',
    'Apply template to new invoice: pre-fills seller, logo, colors, defaults',
    'Platform default templates bundled (PLATFORM_DEFAULT_TEMPLATE + PLATFORM_LETTER_TEMPLATE)',
    'TemplateDesignLayout: renders A4 canvas with layout elements, zoom, snap-to-grid toggle',
]:
    bullet(item, colour=GREEN)

h3('Gap A — Template Designer Drag-Drop is Non-Functional', RED)
lv('File:','src/components/invoice/TemplateDesignLayout.tsx (lines 26–110, 316–322)')
para('The state variables and pointer handlers exist (draggingId, handlePointerDown/Move/Up), '
     'but elements do not show visual feedback during drag (no cursor change, no ghost, no opacity). '
     'Boundary constraints are applied only on pointer-up, not during move, making the '
     'experience feel broken. Elements of types "items", "notes", "footer" are excluded '
     'from dragging with no visual indicator.')
para('Proposed solution:', bold=True)
code(
"""// TemplateDesignLayout.tsx — update element JSX render (near line 316):
<div
  key={el.id}
  data-id={el.id}
  onPointerDown={canDrag ? e => handlePointerDown(e, el.id) : undefined}
  style={{
    position: 'absolute',
    left: el.x, top: el.y, width: el.width,
    cursor: draggingId === el.id ? 'grabbing' : canDrag ? 'grab' : 'default',
    opacity: draggingId === el.id ? 0.70 : 1,
    outline: selectedId === el.id ? '2px solid #6B21A8' : 'none',
    boxShadow: draggingId === el.id ? '0 4px 12px rgba(107,33,168,0.3)' : 'none',
    zIndex: draggingId === el.id ? 999 : 1,
    userSelect: 'none',
    transition: draggingId === el.id ? 'none' : 'opacity 0.15s, box-shadow 0.15s',
  }}
>

// Update handlePointerMove to enforce boundaries DURING drag (lines 86-106):
const handlePointerMove = (e: React.PointerEvent) => {
  if (!draggingId || !dragStartPos || !elementStartPos) return;
  const dx = (e.clientX - dragStartPos.x) / zoom;
  const dy = (e.clientY - dragStartPos.y) / zoom;
  let nx = elementStartPos.x + dx;
  let ny = elementStartPos.y + dy;
  if (snapToGrid) { nx = Math.round(nx/GRID)*GRID; ny = Math.round(ny/GRID)*GRID; }
  // Clamp to canvas boundaries immediately on every move event
  const el = layout.find(l => l.id === draggingId);
  nx = Math.max(0, Math.min(nx, A4_W - (el?.width ?? 40)));
  ny = Math.max(0, Math.min(ny, A4_H - 20));
  setLayout(prev => prev.map(l => l.id === draggingId ? {...l, x: nx, y: ny} : l));
};

// For non-draggable elements, show a lock badge:
{!canDrag && (
  <div style={{position:'absolute',top:2,right:2,opacity:0.4}}>
    <Lock size={10} />
  </div>
)}
""")
lv('Effort:','2–3 hours (pure frontend)',vc=BLUE)

h3('Gap B — No Live Preview in Template Cards', AMBER)
para('TemplateLibrary shows cards with name and description only. No thumbnail of how '
     'an invoice would look using that template.')
para('Proposed solution: render a miniature (scaled-down) invoice canvas inside each card '
     'using the template\'s colors and logo, or generate a thumbnail on template save.')
lv('Effort:','3–4 hours',vc=BLUE)

# ── 3.2 Company Settings ─────────────────────────────────────────
h2('3.2  Company Settings')
lv('Files:','src/components/screens/Settings.tsx · api/app/Controllers/CompanyProfileController.php')
lv('Status:','PARTIAL — 6/10',vc=AMBER)

h3('What Works')
for item in [
    'Company name, VAT ID, address, email, phone, IBAN, BIC',
    'Invoice defaults: currency, tax rate, payment terms text',
    'invoiceNumberFormat and letterNumberFormat fields exist in DB and type',
    'defaultTemplateId stored in profile',
]:
    bullet(item, colour=GREEN)

h3('Gap A — Logo Upload Missing from Settings', RED)
lv('File:','src/components/screens/Settings.tsx (lines 287–319)')
para('Settings.tsx shows a plain text input for logoUrl. TemplateEditor.tsx has a complete '
     'base64 FileReader upload (lines 106–141). The logic needs to be ported to Settings.tsx.')
code(
"""// Settings.tsx — add near top of component:
const logoRef = useRef<HTMLInputElement>(null);

const handleLogoUpload = (e: React.ChangeEvent<HTMLInputElement>) => {
  const file = e.target.files?.[0]; if(!file) return;
  if(!file.type.startsWith('image/')) { toast.error('Image files only'); return; }
  if(file.size > 2*1024*1024) { toast.error('Max 2 MB'); return; }
  const reader = new FileReader();
  reader.onloadend = () => handleChange('logoUrl', reader.result as string);
  reader.readAsDataURL(file);
};

// Replace the URL input (line 291) with:
<input ref={logoRef} type="file" accept="image/*" onChange={handleLogoUpload} className="hidden"/>
<div className="flex items-center gap-3">
  {editedProfile.logoUrl && (
    <img src={editedProfile.logoUrl} className="h-14 w-auto rounded border object-contain"/>
  )}
  <Button variant="outline" size="sm" onClick={() => logoRef.current?.click()}>
    <Upload className="h-4 w-4 mr-2"/> Upload Logo
  </Button>
</div>
""")
lv('Effort:','1 hour',vc=BLUE)

h3('Gap B — Invoice Number Format Has No Builder or Preview', AMBER)
lv('File:','src/components/screens/Settings.tsx (lines 405–418)')
para('Users see a plain text input but have no way to preview the result or discover '
     'valid tokens ({YYYY}, {YY}, {NNN…}, {MM}).')
code(
"""// Settings.tsx — add preview helper:
const previewFmt = (fmt:string) => fmt
  .replace('{YYYY}', String(new Date().getFullYear()))
  .replace('{YY}',   String(new Date().getFullYear()).slice(-2))
  .replace('{MM}',   String(new Date().getMonth()+1).padStart(2,'0'))
  .replace(/{N+}/g, m => '1'.padStart(m.length-2,'0'));

// Replace input block:
<div className="flex gap-2 items-center">
  <Input
    value={editedProfile.invoiceNumberFormat || ''}
    onChange={e => handleChange('invoiceNumberFormat', e.target.value)}
    placeholder="INV-{YYYY}-{NNNNN}"
    className="font-mono flex-1"
  />
  <span className="text-sm font-mono px-3 py-2 bg-slate-100 rounded border min-w-[140px]">
    {previewFmt(editedProfile.invoiceNumberFormat || 'INV-{YYYY}-{NNNNN}')}
  </span>
</div>
<div className="flex gap-2 mt-1 flex-wrap">
  {['INV-{YYYY}-{NNNNN}','{YYYY}/{NNN}','{YY}{NNNNN}'].map(p => (
    <Button key={p} variant="outline" size="xs"
      onClick={() => handleChange('invoiceNumberFormat', p)}>{p}</Button>
  ))}
</div>
<p className="text-xs text-slate-400 mt-1">
  Tokens: {'{YYYY} {YY} {MM} {N} {NNN} {NNNNN}'}
</p>
""")
lv('Effort:','1.5 hours',vc=BLUE)

h3('Gap C — Signature Upload Missing', AMBER)
para('A company signature image (for PDF footer) is not exposed in Settings.tsx. '
     'The same FileReader pattern from logo upload can be reused for a signatureUrl field.')
lv('Effort:','1 hour',vc=BLUE)

doc.add_page_break()

# ── 3.3 Workspace ────────────────────────────────────────────────
h2('3.3  Workspace / File Manager')
lv('Files:','src/components/screens/Workspace.tsx · api/app/Controllers/WorkspaceController.php')
lv('Status:','PARTIAL — 7/10 (functional but has a critical security gap)',vc=AMBER)

h3('What Works')
for item in [
    'List files and folders with name, size, type, modified date',
    'Upload files (multi-file, progress bar, 100 MB limit per file)',
    'Create folders (mkdir)',
    'Delete files/folders (single and bulk)',
    'Download single file and download multiple as ZIP',
    'Rename files and folders (updates all child paths for directories)',
    'Open file locally using OS command (Windows/Mac/Linux)',
    'Standard text search against workspace_files table (name + content columns)',
    'AI natural-language search using Gemini 2.5-flash API',
    'AI query history stored in aiquery_history table',
    'ZIP extraction with options: inline / to new folder / delete source',
    'Content indexing on upload via ContentExtractor',
    'Breadcrumb navigation, bulk selection, sort by name/size/type/date',
]:
    bullet(item, colour=GREEN)

h3('CRITICAL Gap A — AI Search is Vulnerable to SQL Injection', RED)
lv('File:','api/app/Controllers/WorkspaceController.php (line ~756)')
para('The Gemini API returns a WHERE clause string. The controller executes it '
     'directly against the database without sanitization:')
code(
"""// WorkspaceController.php line ~756 — CURRENT (UNSAFE):
$builder->where("($whereClause)", null, false);
// Comment in code: "a proper SQL parser/sanitizer should be used"
// If Gemini returns: "1=1; DROP TABLE workspace_files--" → executed as-is
""")
para('Proposed fix — whitelist-validate the Gemini output before execution:')
code(
"""// WorkspaceController.php — replace the unsafe where() call:
private function sanitizeAIWhereClause(string $clause): string
{
    // Allow only safe column references and comparison operators
    $allowedColumns = ['name','extension','mime_type','size','content','is_dir','path'];
    $colPattern = implode('|', array_map('preg_quote', $allowedColumns));

    // Strip anything that isn't a comparison on an allowed column
    if (!preg_match('/^[\s\w\.\'\"%\(\)=<>!andornotlikeANDORNOTLIKE]+$/i', $clause)) {
        log_message('warning', 'AI search rejected unsafe clause: ' . $clause);
        return '1=0'; // Return no results instead of executing unsafe query
    }

    // Validate column names actually used
    preg_match_all('/\\b(' . $colPattern . ')\\b/i', $clause, $matches);
    $usedCols = array_unique($matches[1]);
    foreach ($usedCols as $col) {
        if (!in_array(strtolower($col), $allowedColumns)) {
            return '1=0';
        }
    }
    return $clause;
}

// Replace line 756:
$safeClause = $this->sanitizeAIWhereClause($whereClause);
$builder->where("($safeClause)", null, false);
""")
lv('Effort:','2 hours · CRITICAL — fix before production',vc=RED)

h3('Gap B — Workspace Not in Sidebar Navigation', AMBER)
lv('File:','src/components/layout/AppSidebar.tsx')
para('Workspace is registered in App.tsx and lazy-loaded, but no nav item exists in the '
     'sidebar. Users can only reach it via direct URL hash #workspace.')
code(
"""// AppSidebar.tsx — add Workspace nav item in the main nav list:
{ label: t('nav.workspace') || 'Workspace', icon: FolderOpen, screen: 'workspace' },

// en.ts / de.ts / ar.ts / pl.ts — add:
workspace: 'Workspace',    // en
workspace: 'Arbeitsbereich', // de
workspace: 'مساحة العمل',  // ar
workspace: 'Przestrzeń robocza', // pl
""")
lv('Effort:','30 minutes',vc=BLUE)

h3('Gap C — Temp ZIP Files Not Cleaned Up', AMBER)
lv('File:','api/app/Controllers/WorkspaceController.php (downloadZip, line ~631)')
para('The controller creates a temp ZIP file and streams it to the client but relies '
     'on PHP garbage collection for cleanup. On high-traffic servers temp files accumulate.')
code(
"""// WorkspaceController.php — add explicit cleanup in downloadZip():
$tmpFile = tempnam(sys_get_temp_dir(), 'ws_zip_');
// ... create zip ...
$response = $this->response->download($tmpFile, null);
register_shutdown_function(function() use ($tmpFile) {
    if (file_exists($tmpFile)) unlink($tmpFile);
});
return $response;
""")
lv('Effort:','30 minutes',vc=BLUE)

h3('Gap D — Content Indexing Failures Are Silent', AMBER)
para('ContentExtractor::extract() failures are caught silently (lines ~419–421). '
     'Files with unindexable content (encrypted PDFs, binary) have empty content '
     'columns and become unsearchable with no user notification.')
lv('Effort:','1 hour — add user-visible warning badge on non-indexed files',vc=BLUE)

doc.add_page_break()

# ── 3.4 Customer Billing / Subscription ──────────────────────────
h2('3.4  Customer Billing & Subscription Management')
lv('Files:','src/components/screens/Billing.tsx · api/app/Controllers/Billing.php · src/services/api.ts (billingService)')
lv('Status:','PARTIAL — 6/10',vc=AMBER)

h3('What Works')
for item in [
    'GET /billing/subscription → current plan name, price, limits, usage stats',
    'GET /billing/plans → all active public plans with features',
    'POST /billing/upgrade → creates Stripe Checkout session, redirects to payment page',
    'GET /billing/package-services → service definitions',
    'Billing.tsx: shows current plan, usage bar, plan comparison, upgrade button',
    'PackageComparison.tsx: feature comparison table across all plans',
]:
    bullet(item, colour=GREEN)

h3('Gap A — Payment History Returns Empty Array (STUB)', RED)
lv('File:','api/app/Controllers/Billing.php (lines 158–163)')
code(
"""// Billing.php — CURRENT STUB:
public function history()
{
    // TODO: Sync invoices from Stripe via Webhook to local DB
    return $this->response->setJSON([])->setStatusCode(200);
}
""")
para('Proposed fix — implement Stripe webhook endpoint and local invoice storage:')
code(
"""// Step 1: Create Stripe webhook handler (new file):
// api/app/Controllers/StripeWebhook.php
public function handle()
{
    $payload = $this->request->getBody();
    $sig     = $this->request->getHeaderLine('Stripe-Signature');
    $secret  = env('STRIPE_WEBHOOK_SECRET');

    try {
        $event = \\Stripe\\Webhook::constructEvent($payload, $sig, $secret);
    } catch (\\Exception $e) {
        return $this->fail($e->getMessage(), 400);
    }

    if ($event->type === 'invoice.payment_succeeded') {
        $inv  = $event->data->object;
        $sub  = $inv->subscription;
        // Save to billing_invoices table
        $model = new \\App\\Models\\BillingInvoiceModel();
        $model->upsert([
            'stripe_invoice_id'  => $inv->id,
            'stripe_customer_id' => $inv->customer,
            'amount'             => $inv->amount_paid / 100,
            'currency'           => strtoupper($inv->currency),
            'status'             => 'paid',
            'invoice_pdf_url'    => $inv->invoice_pdf,
            'period_start'       => date('Y-m-d', $inv->period_start),
            'period_end'         => date('Y-m-d', $inv->period_end),
            'paid_at'            => date('Y-m-d H:i:s', $inv->status_transitions->paid_at),
        ]);
    }
    return $this->respond(['received' => true]);
}

// Step 2: Update history() to query billing_invoices table:
public function history()
{
    $tenantId = session()->get('tenantId');
    // Get stripe_customer_id for this tenant
    $sub      = (new SubscriptionModel())->where('tenant_id', $tenantId)->first();
    if (!$sub) return $this->respond(['data' => []]);

    $invoices = (new BillingInvoiceModel())
        ->where('stripe_customer_id', $sub['stripe_customer_id'])
        ->orderBy('paid_at','DESC')
        ->findAll();

    return $this->respond(['data' => $invoices]);
}

// Step 3: Add migration for billing_invoices table:
// 2026-XX-XX_CreateBillingInvoicesTable.php
""")
lv('Effort:','5–6 hours (webhook + migration + history endpoint + frontend)',vc=BLUE)

h3('Gap B — Stripe Price IDs are Hardcoded Mock Values', RED)
lv('File:','api/app/Controllers/Billing.php (line 134)')
code(
"""// CURRENT — hardcoded fake IDs:
$priceId = ($newPlanId == 2) ? 'price_pro_monthly' : 'price_starter_monthly';

// FIX — store stripe_price_id in plans table:
// Migration: ALTER TABLE plans ADD COLUMN stripe_price_id VARCHAR(64) NULL;
// Then in upgrade():
$plan = (new PlanModel())->find($data['plan_id']);
if (empty($plan['stripe_price_id'])) {
    return $this->fail('Payment not configured for this plan', 422);
}
$session = \\Stripe\\Checkout\\Session::create([
    'mode'       => 'subscription',
    'line_items' => [['price' => $plan['stripe_price_id'], 'quantity' => 1]],
    'success_url' => env('APP_URL') . '/billing?session_id={CHECKOUT_SESSION_ID}',
    'cancel_url'  => env('APP_URL') . '/billing',
    'customer'    => $stripeCustomerId,
]);
""")
lv('Effort:','2 hours',vc=BLUE)

h3('Gap C — No Plan Downgrade or Cancellation', AMBER)
para('Billing.php and Billing.tsx only support upgrading. No endpoint exists for '
     'downgrading a plan or cancelling a subscription.')
code(
"""// New endpoint: POST /billing/cancel
public function cancel()
{
    $sub = (new SubscriptionModel())->where('tenant_id', $tenantId)->first();
    \\Stripe\\Subscription::update($sub['stripe_subscription_id'], [
        'cancel_at_period_end' => true,
    ]);
    (new SubscriptionModel())->update($sub['id'], ['cancel_at_period_end' => 1]);
    return $this->respond(['message' => 'Subscription will cancel at period end']);
}
""")
lv('Effort:','2 hours',vc=BLUE)

h3('Gap D — Feature/Limit Enforcement Missing Entirely', RED)
para('The DB has plans.limits JSON and tenant_usage records, but NO controller checks '
     'usage before allowing resource creation. Users can exceed plan limits.')
code(
"""// Create api/app/Traits/UsageEnforcementTrait.php:
trait UsageEnforcementTrait
{
    protected function checkLimit(string $resource): bool
    {
        $tenantId = session()->get('tenantId');
        $sub      = (new SubscriptionModel())->where('tenant_id',$tenantId)->first();
        $plan     = (new PlanModel())->find($sub['plan_id']);
        $limits   = json_decode($plan['limits'] ?? '{}', true);
        $limit    = $limits[$resource] ?? -1;   // -1 = unlimited

        if ($limit === -1 || $limit === 0) return true;  // unlimited

        $usage = (new TenantUsageModel())
            ->where('tenant_id', $tenantId)
            ->where('resource_key', $resource)
            ->first();
        $used = $usage['used_amount'] ?? 0;

        return $used < $limit;
    }
}

// Use in InvoiceController::create():
use UsageEnforcementTrait;

public function create()
{
    if (!$this->checkLimit('invoices')) {
        return $this->fail('Invoice limit reached for your plan. Please upgrade.', 429);
    }
    // ... rest of create logic
    // After successful create, increment usage:
    (new TenantUsageModel())->increment($tenantId, 'invoices');
}
""")
lv('Effort:','4–5 hours (trait + hooks in all relevant controllers)',vc=BLUE)

doc.add_page_break()

# ── 3.5 Authentication ────────────────────────────────────────────
h2('3.5  Authentication')
lv('Files:','api/app/Controllers/Auth.php · src/components/screens/Login.tsx · src/components/screens/Signup.tsx · src/stores/authStore.ts')
lv('Status:','PARTIAL — 7/10',vc=AMBER)

h3('What Works')
for item in [
    'Login with email/password → JWT issued',
    'Signup with tenant creation (company name, subdomain auto-generated)',
    'Logout clears Zustand authStore and localStorage',
    'Token refresh endpoint',
    'Forgot password → email link → reset password flow',
    'Auto-logout on 401 via Axios response interceptor (api.ts:41–55)',
    'Token migration from old localStorage key format',
    'URL parameter token handling for quick-access links',
]:
    bullet(item, colour=GREEN)

h3('Gap A — No Two-Factor Authentication', RED)
para('Auth.php issues a JWT immediately after password validation. No TOTP step exists. '
     'The users table has no totp_secret or totp_enabled column.')
code(
"""// Complete 2FA flow — see detailed implementation in the companion gap document.
// Summary of changes needed:
// 1. DB: ALTER TABLE users ADD totp_secret VARCHAR(64), ADD totp_enabled TINYINT(1) DEFAULT 0
// 2. Backend: composer require spomky-labs/otphp
// 3. Auth.php login(): if totp_enabled → return {requiresTwoFactor, interimToken} instead of JWT
// 4. New endpoint POST /auth/verify-totp: validate code, return real JWT
// 5. New endpoint POST /auth/setup-totp: generate QR URI for authenticator app
// 6. Login.tsx: add two-step UI (credentials → TOTP code input)
// 7. Settings.tsx: add 2FA enable/disable section with QR code display
""")
lv('Effort:','5–6 hours',vc=BLUE)

h3('Gap B — customerApi.ts Has No Request Interceptor', AMBER)
lv('File:','src/services/customerApi.ts (lines 46–115)')
para('Every method manually passes Authorization headers. If the token expires, '
     'the 401 interceptor redirects but does NOT attempt token refresh first.')
code(
"""// customerApi.ts — add after axios.create():
import { useAuthStore } from '../stores/authStore';
customerApi.interceptors.request.use((config) => {
  const token = useAuthStore.getState().token;
  if (token) {
    config.headers.Authorization   = 'Bearer ' + token;
    config.headers['X-Authorization'] = 'Bearer ' + token;
  }
  return config;
});
// Then remove all manual { headers: {...} } from every method call.
""")
lv('Effort:','1 hour',vc=BLUE)

h3('Gap C — Password Reset Sends Hardcoded "password123"', RED)
lv('File:','api/app/Controllers/AdminUsers.php (line ~252)')
para('The admin "reset password" action sets the user\'s password to the hardcoded '
     'string "password123". This is a critical security vulnerability.')
code(
"""// AdminUsers.php reset() — CURRENT:
$password = 'password123'; // INSECURE

// FIX:
$password = bin2hex(random_bytes(8)); // e.g. "a3f2e8c1b7d9e4f0"
// Hash before storing:
$hashed = password_hash($password, PASSWORD_BCRYPT);
(new UserModel())->update($userId, [
    'password'               => $hashed,
    'password_reset_required'=> 1,  // Force change on next login
]);
// Send new password via email (not in response):
EmailService::sendPasswordReset($user['email'], $password);
return $this->respond(['message' => 'Password reset. User will receive an email.']);
""")
lv('Effort:','1 hour · CRITICAL security fix',vc=RED)

doc.add_page_break()

# ── 3.6 Quick Access Portal ──────────────────────────────────────
h2('3.6  Quick Access Portal')
lv('Files:','src/components/screens/QuickAccessInvoice.tsx · api/app/Controllers/QuickAccessAuth.php · api/app/Database/Migrations/2026-02-27_CreateQuickAccessSessions.php')
lv('Status:','PARTIAL — 5/10',vc=AMBER)

h3('What Is Quick Access')
para('An unauthenticated invoice creation portal. Visitors can create a draft invoice '
     'and preview/download it without signing up. Gated actions (save to account, send, '
     'export) prompt for email + OTP login. Drafts persist in localStorage and can be '
     'restored cross-device via a token URL parameter.')

h3('What Works')
for item in [
    'Guest invoice editor pre-loaded with demo line items',
    'Draft saved to localStorage (key: qa_draft) on every change',
    'Cross-device restore: GET /auth/quick-access/draft?token= fetches saved draft',
    'OTP flow: POST /auth/check-email → POST /auth/quick-access (send OTP) → POST /auth/quick-access/verify',
    'After OTP login, pending action (save/export/download) is executed automatically',
    'Invoice preview modal with PDF and UBL XML tabs',
    'calculateInvoiceTotals() wired for live subtotal/tax/grand total',
]:
    bullet(item, colour=GREEN)

h3('Gap A — OTP Rate Limiting Missing', RED)
para('POST /auth/quick-access (send OTP) has no rate limiting. An attacker can spam '
     'OTP requests to any email address indefinitely.')
code(
"""// QuickAccessAuth.php — add at top of sendOtp():
$email   = $data['email'] ?? '';
$cacheKey = 'otp_rate_' . md5($email);
$attempts = cache()->get($cacheKey) ?? 0;

if ($attempts >= 3) {
    return $this->fail('Too many OTP requests. Try again in 1 hour.', 429);
}
cache()->save($cacheKey, $attempts + 1, 3600); // 1-hour window
""")
lv('Effort:','30 minutes',vc=BLUE)

h3('Gap B — Draft Token Security', AMBER)
para('The qa_token URL parameter is a plain string stored in quick_access_sessions. '
     'If it is sequential or short, it can be guessed to retrieve other users\' drafts.')
code(
"""// QuickAccessAuth.php — when creating a session token:
$token = bin2hex(random_bytes(32)); // 64-char hex, cryptographically random
// Store with short TTL:
$model->insert(['token' => $token, 'draft' => json_encode($draft), 'expires_at' => date('Y-m-d H:i:s', time() + 900)]);
""")
lv('Effort:','1 hour',vc=BLUE)

h3('Gap C — Pending Action Uses localStorage (Tab Collision)', AMBER)
para('After OTP login, the pending action (what to do after auth) is stored in localStorage '
     'as JSON. If the user has multiple tabs open, tabs can overwrite each other\'s pending action.')
lv('Proposed fix:','Use React state or sessionStorage (tab-scoped) instead of localStorage.',vc=DARK)
lv('Effort:','1 hour',vc=BLUE)

h3('Gap D — No Invoice Sharing Link', AMBER)
para('The architecture supports sharing via token, but no endpoint exists to retrieve '
     'a shared invoice by token. The frontend references this concept but the backend '
     'endpoint is missing.')
lv('Effort:','2 hours — add GET /public/invoices/:shareToken endpoint',vc=BLUE)

doc.add_page_break()

# ── 3.7 Admin — Packages ─────────────────────────────────────────
h2('3.7  Admin — Packages & Subscription Plans')
lv('Files:','src/components/screens/Admin/SApackages.tsx · SAPackageForm.tsx · SAPackageServices.tsx · api/app/Controllers/AdminPackages.php · AdminPackageServices.php')
lv('Status:','PARTIAL — 6/10',vc=AMBER)

h3('What Works')
for item in [
    'SApackages.tsx: lists all plans with inline editing of name, price, currency, duration, status',
    'Service columns dynamically rendered from adminPackageServicesService.getAll()',
    'Feature switches and numeric inputs per package × service',
    'Delete package (blocked if trailing or used by tenants)',
    'Set trailing/default package (only one allowed)',
    'SAPackageServices.tsx: full CRUD for service definitions',
    'Service types: storage, users, api_calls, bandwidth, invoices, projects, custom',
    'AdminPackages.php: create/update converts features JSON → limits JSON via syncLimitsFromFeatures()',
]:
    bullet(item, colour=GREEN)

h3('Gap A — Currency is Hardcoded to EUR in Backend', AMBER)
lv('File:','api/app/Controllers/AdminPackages.php (line ~40)')
code(
"""// AdminPackages.php line 40 — CURRENT:
'currency' => 'EUR',   // Hardcoded regardless of plan setting

// FIX — store currency in plans table:
// Migration: ALTER TABLE plans ADD COLUMN currency CHAR(3) DEFAULT 'EUR';
// Then use: 'currency' => $plan['currency'] ?? 'EUR',
""")
lv('Effort:','1 hour',vc=BLUE)

h3('Gap B — Features JSON Has No Schema Validation', AMBER)
para('Features can be any shape. If a service type is removed but old feature data '
     'references it, syncLimitsFromFeatures() silently skips it. No type enforcement.')
lv('Proposed fix:','Validate incoming features against current package_services definitions before saving.',vc=DARK)
lv('Effort:','1.5 hours',vc=BLUE)

h3('Gap C — No Plan Retirement Workflow', AMBER)
para('If a plan is deactivated (is_active=0), existing subscribers are not notified '
     'and the plan still appears in their subscription details. No migration path exists.')
lv('Effort:','3 hours — add notification email + grace period logic',vc=BLUE)

h3('Gap D — Upgrade Plan Endpoint Is a Stub', RED)
lv('File:','api/app/Controllers/AdminUsers.php (lines 213–219)')
code(
"""// AdminUsers.php upgradePlan() — CURRENT STUB:
public function upgradePlan($id = null)
{
    return $this->respond(['success' => true, 'message' => 'Plan upgraded successfully']);
    // Does nothing — subscription unchanged
}

// FIX — implement actual plan change:
public function upgradePlan($id = null)
{
    $data   = $this->request->getJSON(true);
    $planId = $data['planId'] ?? null;
    if (!$planId) return $this->fail('planId is required');

    $sub  = (new SubscriptionModel())->where('tenant_id', $id)->first();
    $plan = (new PlanModel())->find($planId);
    if (!$plan) return $this->failNotFound('Plan not found');

    // Update Stripe subscription if stripe_subscription_id exists:
    if (!empty($sub['stripe_subscription_id']) && !empty($plan['stripe_price_id'])) {
        $stripeSub = \\Stripe\\Subscription::retrieve($sub['stripe_subscription_id']);
        \\Stripe\\Subscription::update($sub['stripe_subscription_id'], [
            'items' => [['id' => $stripeSub->items->data[0]->id, 'price' => $plan['stripe_price_id']]],
            'proration_behavior' => 'create_prorations',
        ]);
    }

    (new SubscriptionModel())->update($sub['id'], ['plan_id' => $planId, 'status' => 'active']);
    (new TenantModel())->update($id, ['plan_id' => $planId]);
    $this->logAction('plan_upgraded', '', "Tenant {$id} moved to plan {$plan['name']}");
    return $this->respond(['success' => true]);
}
""")
lv('Effort:','2–3 hours',vc=BLUE)

doc.add_page_break()

# ── 3.8 Admin — Users ────────────────────────────────────────────
h2('3.8  Admin — SaaS User Management')
lv('Files:','src/components/screens/Admin/SAASusers.tsx · SAUserDetails.tsx · api/app/Controllers/AdminUsers.php')
lv('Status:','PARTIAL — 5/10',vc=AMBER)

h3('What Works')
for item in [
    'Paginated list of tenants with search (company name, subdomain, email) and status filter',
    'Tenant detail view: plan, status, usage stats, join date',
    'Suspend tenant: sets status = "suspended"',
    'Activate tenant: sets status = "active", moves trialing → active subscription',
    'Export tenants as CSV (6 columns)',
    'View usage stats: storage, API calls, bandwidth, active users',
]:
    bullet(item, colour=GREEN)

h3('Gap A — Last Login Always Returns Current Time', AMBER)
lv('File:','api/app/Controllers/AdminUsers.php (lines ~83, ~140)')
code(
"""// CURRENT — always shows "now":
'lastLogin' => date('Y-m-d\\TH:i:s\\Z'),

// FIX — query actual last login from users table:
// Migration: ALTER TABLE users ADD COLUMN last_login_at DATETIME NULL;
// Auth.php login() — after successful login, set: last_login_at = NOW()
// AdminUsers.php show():
$primaryUser = (new UserModel())
    ->where('tenant_id', $tenant['id'])
    ->where('role', 'owner')
    ->first();
'lastLogin' => $primaryUser['last_login_at'] ?? null,
""")
lv('Effort:','1.5 hours',vc=BLUE)

h3('Gap B — Default Usage Limits Hardcoded', AMBER)
lv('File:','api/app/Controllers/AdminUsers.php (lines ~300–305)')
code(
"""// CURRENT — hardcoded regardless of plan:
'storageLimit'   => 10,      // GB
'usersLimit'     => 5,
'bandwidthLimit' => 100,     // GB
'apiCallsLimit'  => 50000,

// FIX — read from plan limits:
$plan   = (new PlanModel())->find($subscription['plan_id']);
$limits = json_decode($plan['limits'] ?? '{}', true);
'storageLimit'   => $limits['storage_gb']  ?? 10,
'usersLimit'     => $limits['users']       ?? 5,
'bandwidthLimit' => $limits['bandwidth_gb']?? 100,
'apiCallsLimit'  => $limits['api_calls']   ?? 50000,
""")
lv('Effort:','1 hour',vc=BLUE)

h3('Gap C — Email Fallback Uses Fake Domain', AMBER)
lv('File:','api/app/Controllers/AdminUsers.php (lines ~77, ~134)')
code(
"""// CURRENT — fallback to fake email:
'email' => $user['email'] ?? ($tenant['subdomain'] . '@tech-portal.io'),

// FIX — query actual user email, return null if not found:
$primaryUser = (new UserModel())
    ->select('email')
    ->where('tenant_id', $tenant['id'])
    ->first();
'email' => $primaryUser['email'] ?? null,
""")
lv('Effort:','30 minutes',vc=BLUE)

h3('Gap D — CSV Export Lacks Usage Columns', AMBER)
para('Current export (lines 271–283) only has: ID, company name, subdomain, plan, status, joined. '
     'No usage data (storage, API calls) which would be useful for billing reviews.')
lv('Effort:','1 hour',vc=BLUE)

doc.add_page_break()

# ── 3.9 Admin — Billing ──────────────────────────────────────────
h2('3.9  Admin — Billing & Revenue Dashboard')
lv('Files:','src/components/screens/Admin/SAbilling.tsx · api/app/Controllers/AdminBilling.php')
lv('Status:','PARTIAL — 4/10  (most backend methods return stubs or approximations)',vc=AMBER)

h3('What Works')
for item in [
    'Revenue summary: total revenue, paid vs pending invoice counts',
    'Monthly revenue for last 6 months (estimated from subscription prices)',
    'Create admin invoice: stores subscription record, returns total',
    'List invoices: queries subscriptions table, maps to invoice-like format',
]:
    bullet(item, colour=GREEN)

h3('Gap A — Invoice by ID Returns Hardcoded Mock', RED)
lv('File:','api/app/Controllers/AdminBilling.php (lines 104–124)')
code(
"""// CURRENT — ignores $id, returns fake data:
public function show($id = null)
{
    $invoice = ['id' => $id, 'company' => 'Mock Company', ...]; // Hardcoded!
    return $this->respond(['data' => $invoice]);
}

// FIX — query actual record:
public function show($id = null)
{
    $sub = (new SubscriptionModel())->find($id);
    if (!$sub) return $this->failNotFound();
    $tenant = (new TenantModel())->find($sub['tenant_id']);
    $plan   = (new PlanModel())->find($sub['plan_id']);
    return $this->respond(['data' => [
        'id'      => $sub['id'],
        'company' => $tenant['company_name'],
        'plan'    => $plan['name'],
        'amount'  => $plan['price'],
        'status'  => $sub['status'],
        'period'  => [$sub['current_period_start'], $sub['current_period_end']],
    ]]);
}
""")
lv('Effort:','1 hour',vc=BLUE)

h3('Gap B — Invoice PDF Download Returns Mock Content', RED)
lv('File:','api/app/Controllers/AdminBilling.php (lines 130–139)')
code(
"""// CURRENT — fake PDF bytes:
public function downloadPdf($id = null)
{
    $pdfContent = '%PDF-1.4 Mock PDF Content';  // Not a real PDF
    return $this->response->setHeader('Content-Type','application/pdf')->setBody($pdfContent);
}

// FIX — generate real PDF using TCPDF:
public function downloadPdf($id = null)
{
    $sub    = (new SubscriptionModel())->find($id);
    $tenant = (new TenantModel())->find($sub['tenant_id']);
    $plan   = (new PlanModel())->find($sub['plan_id']);

    $pdf = new \\TCPDF(); // Or use mPDF
    $pdf->AddPage();
    $pdf->writeHTML(view('pdf/subscription_invoice', ['sub'=>$sub,'tenant'=>$tenant,'plan'=>$plan]));
    return $this->response
        ->setHeader('Content-Type','application/pdf')
        ->setHeader('Content-Disposition','attachment; filename="invoice-'.$id.'.pdf"')
        ->setBody($pdf->Output('','S'));
}
""")
lv('Effort:','3 hours',vc=BLUE)

h3('Gap C — Revenue Uses Subscriptions Table Instead of Billing Invoices', RED)
para('AdminBilling.php lines 28–31 queries the subscriptions table and treats each '
     'subscription as an "invoice". This means refunds, credits, partial payments, and '
     'failed payments are all invisible. Revenue figures are theoretical, not actual.')
para('Proper fix requires the Stripe webhook integration described in Gap 3.4-A above '
     'to populate a real billing_invoices table.')
lv('Effort:','Part of billing history fix — 5–6 hours combined',vc=BLUE)

# ── 3.10 Admin — Analytics ───────────────────────────────────────
h2('3.10  Admin — Analytics & Usage Tracking')
lv('Files:','src/components/screens/Admin/SAusage.tsx · api/app/Controllers/AdminAnalytics.php')
lv('Status:','PARTIAL — 5/10',vc=AMBER)

h3('What Works')
for item in [
    'Dashboard stats: total users, active subscriptions, new users this month, churn rate, ARPU',
    'Trend indicators vs previous month (% change)',
    'Revenue history chart (6 months)',
    'User growth history chart (6 months)',
    'Usage metrics by period: storage, API calls, bandwidth, sessions',
    'Per-tenant usage breakdown (sorted by storage)',
]:
    bullet(item, colour=GREEN)

h3('Gap A — Export CSV is Hardcoded Mock', RED)
lv('File:','api/app/Controllers/AdminAnalytics.php (lines 466–482)')
code(
"""// CURRENT — fake CSV:
public function exportUsage()
{
    $csvData = "tenant_id,storage_gb,api_calls\\n1,5.2,1250\\n2,2.8,890\\n";
    return $this->response->setHeader('Content-Type','text/csv')->setBody($csvData);
}

// FIX — actually query and format:
public function exportUsage()
{
    $tenants = (new TenantUsageModel())->findAll();
    $lines   = ["tenant_id,company_name,storage_gb,api_calls,bandwidth_gb,active_users"];
    foreach ($tenants as $row) {
        $tenant   = (new TenantModel())->find($row['tenant_id']);
        $lines[]  = implode(',', [
            $row['tenant_id'],
            '"' . ($tenant['company_name'] ?? '') . '"',
            round(($this->usageModel->getUsage($row['tenant_id'],'storage')['used_amount']??0)/1024/1024/1024, 2),
            $this->usageModel->getUsage($row['tenant_id'],'api_calls')['used_amount'] ?? 0,
            round(($this->usageModel->getUsage($row['tenant_id'],'bandwidth')['used_amount']??0)/1024/1024/1024, 2),
            $this->usageModel->getUsage($row['tenant_id'],'users')['used_amount'] ?? 0,
        ]);
    }
    $csv = implode("\\n", $lines);
    return $this->response
        ->setHeader('Content-Type','text/csv')
        ->setHeader('Content-Disposition','attachment; filename="usage-export.csv"')
        ->setBody($csv);
}
""")
lv('Effort:','1.5 hours',vc=BLUE)

h3('Gap B — Session Count Uses rand() Filler', RED)
lv('File:','api/app/Controllers/AdminAnalytics.php (line ~381)')
code(
"""// CURRENT — random data:
'sessions' => rand(2, 10),  // Fabricated!

// FIX — query actual session count:
// Option A: use ci_sessions table with tenant_id column (requires Auth.php to store tenant_id in session)
// Option B: track session creation in a sessions_log table
// Minimum fix: return null instead of random number:
'sessions' => null,  // Frontend should show "N/A" for unavailable metrics
""")
lv('Effort:','2 hours (proper solution requires session tracking)',vc=BLUE)

h3('Gap C — Bandwidth Calculated from File Size, Not Downloads', AMBER)
para('Bandwidth usage is summed from workspace_files.size (disk usage, not transfer). '
     'True bandwidth = bytes actually downloaded/streamed. These are very different metrics.')
lv('Effort:','3 hours — add download event tracking in WorkspaceController',vc=BLUE)

h3('Gap D — Churn Rate Definition Is Wrong', AMBER)
lv('File:','api/app/Controllers/AdminAnalytics.php (line ~68)')
code(
"""// CURRENT — uses suspended/total (wrong definition of churn):
$churnRate = ($suspendedCount / max($totalUsers, 1)) * 100;

// FIX — churn = subscriptions cancelled THIS PERIOD / subscriptions at START of period:
$cancelledThisMonth = (new SubscriptionModel())
    ->where('status','cancelled')
    ->where('updated_at >=', date('Y-m-01'))
    ->countAllResults();
$activeAtStart = (new SubscriptionModel())
    ->where('created_at <', date('Y-m-01'))
    ->where('status !=','cancelled')
    ->countAllResults();
$churnRate = $activeAtStart > 0 ? ($cancelledThisMonth / $activeAtStart) * 100 : 0;
""")
lv('Effort:','1 hour',vc=BLUE)

doc.add_page_break()

# ── 3.11 Admin — Settings ────────────────────────────────────────
h2('3.11  Admin — System Settings')
lv('Files:','src/components/screens/Admin/SAsettings.tsx · api/app/Controllers/AdminSettings.php')
lv('Status:','PARTIAL — 6/10',vc=AMBER)

h3('What Works')
for item in [
    'Admin profile: name, email',
    'Admin password change',
    'API key management (create, list, revoke)',
    'System settings: maintenance mode, registration control, email config',
    'Database: run migrations, seed data buttons',
]:
    bullet(item, colour=GREEN)
para('Gaps: no email test button, no backup/restore, no system health check endpoint, '
     'no log viewer, no feature flag management.',colour=AMBER)

# ── 3.12 Admin — Tickets ─────────────────────────────────────────
h2('3.12  Admin — Support Tickets')
lv('Files:','src/components/screens/Admin/SATickets.tsx · SATicketDetails.tsx · api/app/Controllers/TicketController.php')
lv('Status:','PARTIAL — 6/10',vc=AMBER)

h3('What Works')
for item in [
    'Ticket creation by customers (TicketController.php create)',
    'List tickets with status filter (open/in_progress/resolved/closed)',
    'View ticket details with full message thread',
    'Admin can update ticket status and reply',
    'Ticket tracking: status change history',
]:
    bullet(item, colour=GREEN)

para('Gaps:')
for item in [
    'No email notification when ticket is created or replied to',
    'No ticket assignment to specific admin users',
    'No priority levels (urgent/normal/low) enforcement',
    'No SLA/response time tracking',
    'No bulk close/resolve action',
    'Ticket creation from customer side (frontend) not verified',
]:
    bullet(item, colour=AMBER)

# ── 3.13 Admin — Wiki ────────────────────────────────────────────
h2('3.13  Admin — Wiki / Knowledge Base')
lv('Files:','src/components/screens/Admin/SAWiki.tsx · api/app/Controllers/Admin/AdminWiki.php')
lv('Status:','PARTIAL — 5/10  (read-only viewer; no edit capability)',vc=AMBER)

h3('What Works')
for item in [
    'File tree sidebar loaded from adminWikiService.getTree()',
    'Markdown content rendered with syntax highlighting',
    'Language selector (wiki supports multiple language folders)',
    'Export current page as PDF (client-side)',
    'Breadcrumb navigation through wiki sections',
]:
    bullet(item, colour=GREEN)

h3('STARTED AND STOPPED: Edit Capability Was Never Built', RED)
para('The backend has the wiki files on disk and AdminWiki.php can read them. '
     'A PUT /admin/wiki/content endpoint does NOT exist in Routes.php. '
     'SAWiki.tsx has no edit state, no textarea, no save button. '
     'The feature was designed (getTree + getContent implemented) but edit was never started.')
code(
"""// Routes.php — add:
$routes->put('admin/wiki/content', 'Admin\\AdminWiki::updateContent');

// AdminWiki.php — add method:
public function updateContent()
{
    $data    = $this->request->getJSON(true);
    $path    = ltrim($data['path'] ?? '', '/');
    $content = $data['content'] ?? '';
    $lang    = $data['language'] ?? 'en';
    $base    = ROOTPATH . 'docs/' . $lang . '/';
    $full    = realpath($base . $path);
    if (!$full || !str_starts_with($full, realpath($base))) {
        return $this->fail('Invalid path', 400);
    }
    file_put_contents($full, $content);
    return $this->respondUpdated(['message' => 'Page saved']);
}

// SAWiki.tsx — add edit mode toggle (see companion gap document for full code)
""")
lv('Effort:','3–4 hours',vc=BLUE)

# ── 3.14 Admin — CMS Pages ───────────────────────────────────────
h2('3.14  Admin — CMS Page Editor (SAPages)')
lv('Files:','src/components/screens/Admin/SAPages.tsx · api/app/Controllers/CmsController.php · api/app/Models/CmsPageModel.php')
lv('Status:','PARTIAL — 6/10',vc=AMBER)

h3('What Works')
for item in [
    'CmsController: GET /admin/cms (list pages), PUT /admin/cms/:slug (update page)',
    'CmsPageModel with slug, title, content (JSON), locale, is_active',
    'Pages seeded: privacy-policy, terms, cookie-policy, impressum (migration 2026-04-18)',
    'Public endpoint: GET /api/public/cms/:slug — no auth required',
]:
    bullet(item, colour=GREEN)

para('Gaps: SAPages.tsx frontend may have limited editor functionality. '
     'Content is stored as JSON sections but the editor UI likely shows a simple textarea. '
     'No draft/publish workflow. No page versioning. No rich-text editor for content blocks.',colour=AMBER)
lv('Effort to complete:','3 hours — wire RichTextEditor per content section, add publish toggle',vc=BLUE)

# ── 3.15 RBAC ────────────────────────────────────────────────────
h2('3.15  Role-Based Access Control (RBAC)')
lv('Files:','src/components/screens/{UserList,UserForm,RoleList,RoleForm}.tsx · api/app/Controllers/{UserController,RoleController,RightController}.php · src/hooks/usePermission.ts')
lv('Status:','PARTIAL — 5/10  (infrastructure exists; enforcement is missing)',vc=AMBER)

h3('What Exists')
for item in [
    'RoleController: CRUD for roles',
    'RightController: list all rights grouped by module',
    'UserController: CRUD for users within a tenant',
    'UserList/UserForm/RoleList/RoleForm screens in admin',
    'usePermission.ts hook with hasPermission() and hasPermissionSync()',
    'Users table has role_id foreign key',
]:
    bullet(item, colour=GREEN)

h3('STARTED AND STOPPED: Permission Checks Are Not Applied', RED)
para('The RBAC infrastructure (roles, rights, usePermission hook) was built but permission '
     'checks are not applied to any screens or API endpoints. Any authenticated user can '
     'access any feature regardless of their role.')
para('The admin portal has zero per-role access control. All admins see all screens.')
code(
"""// Current state — usePermission.ts exists but is never called in screens:
// grep -r "usePermission\|hasPermission" src/components/screens/ → 0 results

// Fix plan:
// 1. Define permission constants (e.g. 'invoice.create', 'template.delete', 'admin.users.view')
// 2. Assign rights to roles via RoleForm
// 3. Include rights array in JWT payload or fetch on login
// 4. Call usePermission() at component level to show/hide UI elements
// 5. Add middleware in backend to validate rights on API endpoints

// Example — InvoiceList.tsx:
const { hasPermission } = usePermission();
{hasPermission('invoice.delete') && (
  <Button onClick={() => handleDelete(invoice.id)}>Delete</Button>
)}

// Example — AdminLayout.tsx (see Gap 5 in companion document for full RBAC guard)
""")
lv('Effort:','6–8 hours (permission constants + hook wiring + backend middleware)',vc=BLUE)

# ── 3.16 AI Invoice Assistant ────────────────────────────────────
h2('3.16  AI Invoice Assistant')
lv('Files:','src/components/screens/AIHistory.tsx · api/app/Controllers/AIInvoiceController.php · src/services/api.ts (aiInvoiceService)')
lv('Status:','PARTIAL — 6/10',vc=AMBER)

h3('What Works')
for item in [
    'POST /ai/parse-invoice: sends user prompt to Gemini API, returns structured invoice JSON',
    'aiInvoiceService.parseInvoicePrompt() in api.ts',
    'AI History screen: lists past AI prompts with search, sort, pagination',
    'WorkspaceController AI search: natural language → SQL WHERE clause via Gemini',
]:
    bullet(item, colour=GREEN)

para('Gaps:')
for item in [
    'AI invoice creation is wired in AIInvoiceController but no "Create from AI" entry point in InvoiceEditor',
    'No UI to trigger AI parsing from the main invoice creation flow',
    'AI History screen exists but "Load into editor" action is unclear',
    'Gemini API key is in backend .env — if missing, AI features fail silently',
    'No rate limiting on AI endpoints (expensive API calls)',
    'WorkspaceController AI search has SQL injection risk (see Gap 3.3-A)',
]:
    bullet(item, colour=AMBER)
lv('Effort to complete AI invoice creation UI:','2–3 hours',vc=BLUE)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 4 — STARTED-AND-STOPPED FEATURES
# ══════════════════════════════════════════════════════════════════
h1('4. Features Started and Then Stopped')
divider()
para('These are features where scaffolding, models, or partial UI were built, '
     'then development was stopped before the feature was functional.')

tbl(
    ['Feature','What Was Built','What Was Stopped','Impact','Effort to Finish'],
    [
        ('Subscription History / Billing Invoices',
         'Billing.tsx UI table, GET /billing/history route',
         'Backend returns empty [] with TODO comment. No Stripe webhook. No DB table for paid invoices.',
         'HIGH — users cannot see past payments',
         '5–6 h'),
        ('Plan Usage Enforcement',
         'TenantUsageModel, tenant_usage DB table, limits in plans.limits JSON',
         'No controller checks limits before allowing resource creation. Zero enforcement.',
         'HIGH — users can exceed plan limits',
         '4–5 h'),
        ('Wiki Admin Edit',
         'SAWiki.tsx read-only viewer, AdminWiki.php getTree+getContent',
         'No PUT endpoint, no edit UI in SAWiki.tsx, no save handler.',
         'MEDIUM',
         '3–4 h'),
        ('Admin RBAC Enforcement',
         'RoleController, RightController, UserList/RoleList screens, usePermission hook',
         'Permission checks never applied to any screen or API endpoint.',
         'HIGH — no access control',
         '6–8 h'),
        ('Invoice Clone / Persist',
         'handleDuplicate() in InvoiceList.tsx, UI button exists',
         'Creates invoice in local React state only. Never calls invoiceService.create(). Lost on refresh.',
         'MEDIUM',
         '1–1.5 h'),
        ('Admin Upgrade Plan',
         'SAUserDetails.tsx upgrade button, POST /admin/users/:id/upgrade route',
         'AdminUsers.php upgradePlan() returns success stub. Subscription unchanged.',
         'HIGH',
         '2–3 h'),
        ('Admin Invoice PDF Download',
         'SAbilling.tsx download button, GET /admin/billing/:id/pdf route',
         'AdminBilling.php downloadPdf() returns "%PDF Mock PDF Content" string.',
         'MEDIUM',
         '3 h'),
        ('Admin Analytics CSV Export',
         'SAusage.tsx export button, GET /admin/analytics/export-usage route',
         'AdminAnalytics.php exportUsage() returns hardcoded 2-row CSV.',
         'MEDIUM',
         '1.5 h'),
        ('Letter Body Rich-Text Editor',
         'invoice.body field typed, isBusinessLetter flag, RichTextEditor component exists',
         'No conditional UI renders a rich-text editor for letter mode in InvoiceEditor.',
         'MEDIUM',
         '1.5–2 h'),
        ('Invoice Sharing via Link',
         'Quick access session tokens, cross-device draft restore',
         'No GET /public/invoices/:shareToken endpoint. No share button in InvoiceList.',
         'LOW–MEDIUM',
         '2 h'),
        ('AI Invoice Entry Point',
         'AIInvoiceController.php fully implemented, AIHistory screen exists',
         'No "Create from AI prompt" button in InvoiceEditor or Dashboard.',
         'LOW',
         '2–3 h'),
        ('Workspace in Sidebar',
         'Workspace.tsx fully built, registered in App.tsx',
         'No nav item in AppSidebar. Users cannot discover the feature.',
         'HIGH (discoverability)',
         '30 min'),
        ('Stripe Webhook Sync',
         'Upgrade creates Stripe session. Stripe price IDs in .env comments.',
         'No webhook endpoint. No billing_invoices table. No event handling.',
         'HIGH',
         '5–6 h'),
        ('2FA / TOTP',
         'Login.tsx has forgotPassword mode showing the pattern. Auth flow is ready.',
         'No TOTP library, no DB columns, no second step in login.',
         'MEDIUM',
         '5–6 h'),
    ],
    widths=[1.6,1.85,1.85,0.8,0.9]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 5 — SECURITY ISSUES
# ══════════════════════════════════════════════════════════════════
h1('5. Security Issues')
divider()

tbl(
    ['Severity','Location','Issue','Fix Summary'],
    [
        ('CRITICAL','WorkspaceController.php line ~756',
         'AI search executes Gemini-generated SQL WHERE clause directly — SQL injection possible',
         'Whitelist allowed columns; validate output before execution'),
        ('CRITICAL','AdminUsers.php line ~252',
         'Password reset hardcodes "password123"',
         'Generate random password; send via email; force change on login'),
        ('HIGH','QuickAccessAuth.php',
         'OTP endpoint has no rate limiting — spam possible',
         'Add 3 attempts per hour per email using CodeIgniter cache'),
        ('HIGH','QuickAccessAuth.php',
         'Session tokens may be short/sequential — guessable',
         'Use bin2hex(random_bytes(32)) for all tokens'),
        ('HIGH','AdminLayout.tsx',
         'All authenticated admins access all screens — no RBAC',
         'Implement permission matrix (see Gap 3.15)'),
        ('HIGH','InvoiceController.php',
         'No status transition validation — any status → any status',
         'Add validateStatusTransition() before update'),
        ('MEDIUM','Billing.php line 144',
         'Checkout success/cancel URLs not origin-validated',
         'Use fixed server-side config URLs, not client-supplied'),
        ('MEDIUM','customerApi.ts',
         'No token refresh on 401 for customer portal',
         'Add request interceptor'),
        ('LOW','CmsController.php',
         'Content type handled inconsistently (array vs JSON string)',
         'Standardize to always decode/encode JSON'),
    ],
    widths=[0.75,2.0,2.5,1.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 6 — FULL BACKLOG (PRIORITISED)
# ══════════════════════════════════════════════════════════════════
h1('6. Complete Prioritised Backlog')
divider()

h2('Sprint 1 — Critical & Security (Week 1, ~18 hours)')
tbl(
    ['#','Task','File(s)','Effort'],
    [
        ('S1-1','FIX AI search SQL injection','WorkspaceController.php ~756','2 h'),
        ('S1-2','FIX hardcoded password reset','AdminUsers.php ~252','1 h'),
        ('S1-3','Implement plan usage enforcement','UsageEnforcementTrait + controllers','4–5 h'),
        ('S1-4','Add Workspace to sidebar navigation','AppSidebar.tsx','0.5 h'),
        ('S1-5','Wire invoice clone to backend','InvoiceList.tsx handleDuplicate','1 h'),
        ('S1-6','Fix Admin upgrade plan stub','AdminUsers.php upgradePlan()','2–3 h'),
        ('S1-7','Add OTP rate limiting','QuickAccessAuth.php','0.5 h'),
        ('S1-8','Fix invoice status transition enforcement','InvoiceController.php','2 h'),
    ],
    widths=[0.6,2.5,2.2,0.65]
)

h2('Sprint 2 — High Priority UX (Week 2, ~20 hours)')
tbl(
    ['#','Task','File(s)','Effort'],
    [
        ('S2-1','Company logo upload in Settings','Settings.tsx','1 h'),
        ('S2-2','Invoice number format builder with preview','Settings.tsx','1.5 h'),
        ('S2-3','Template Designer drag-drop visual feedback','TemplateDesignLayout.tsx','2–3 h'),
        ('S2-4','Letter body rich-text editor','InvoiceEditor.tsx','1.5–2 h'),
        ('S2-5','Stripe billing history (webhook + DB + UI)','Billing.php + StripeWebhook.php','5–6 h'),
        ('S2-6','Stripe price ID from DB (not hardcoded)','Billing.php + plans table migration','2 h'),
        ('S2-7','Admin Analytics real CSV export','AdminAnalytics.php exportUsage()','1.5 h'),
        ('S2-8','Admin Billing: real invoice by ID and PDF','AdminBilling.php show()+downloadPdf()','3 h'),
    ],
    widths=[0.6,2.5,2.2,0.65]
)

h2('Sprint 3 — Medium Priority (Week 3, ~22 hours)')
tbl(
    ['#','Task','File(s)','Effort'],
    [
        ('S3-1','Admin RBAC permission enforcement','AdminLayout.tsx + usePermission hook','6–8 h'),
        ('S3-2','Wiki admin edit UI + endpoint','SAWiki.tsx + AdminWiki.php','3–4 h'),
        ('S3-3','Implement 2FA / TOTP','Auth.php + Login.tsx + DB migration','5–6 h'),
        ('S3-4','Fix last login to use real timestamp','AdminUsers.php + Auth.php + DB migration','1.5 h'),
        ('S3-5','Fix admin analytics session rand() filler','AdminAnalytics.php','2 h'),
        ('S3-6','Fix churn rate calculation','AdminAnalytics.php line ~68','1 h'),
        ('S3-7','customerApi.ts request interceptor','customerApi.ts','1 h'),
    ],
    widths=[0.6,2.5,2.2,0.65]
)

h2('Sprint 4 — Low Priority / Polish (Week 4+, ~22 hours)')
tbl(
    ['#','Task','File(s)','Effort'],
    [
        ('S4-1','Stripe plan cancellation/downgrade','Billing.php','2 h'),
        ('S4-2','Email notifications on invoice events','EmailService.php','4–5 h'),
        ('S4-3','Batch PDF/ZIP export (JSZip)','invoice-export.ts + InvoiceList.tsx','3–4 h'),
        ('S4-4','AI invoice creation entry point in editor','InvoiceEditor.tsx + AIInvoiceController','2–3 h'),
        ('S4-5','Invoice sharing link','App.tsx + new public endpoint','2 h'),
        ('S4-6','Admin invoice PDF with TCPDF','AdminBilling.php','3 h'),
        ('S4-7','Signature upload in Settings','Settings.tsx','1 h'),
        ('S4-8','Workspace temp ZIP cleanup','WorkspaceController.php','0.5 h'),
        ('S4-9','PDF/A-3 with embedded UBL','New backend endpoint + TCPDF','6–8 h'),
        ('S4-10','Usage reset cron job','New scheduled task','1.5 h'),
    ],
    widths=[0.6,2.5,2.2,0.65]
)

# ══════════════════════════════════════════════════════════════════
# SECTION 7 — MASTER STATUS TABLE
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('7. Master Status Table — All Features')
divider()

master = [
    ('Invoice CRUD',                  'DONE',    '10/10','InvoiceEditor + InvoiceController'),
    ('Invoice EN 16931 Validation',   'DONE',    '10/10','invoice-validation.ts'),
    ('Invoice PDF Export',            'DONE',    '10/10','invoice-pdf.ts (client-side)'),
    ('Invoice UBL/JSON/CSV Export',   'DONE',    '9/10', 'invoice-export.ts'),
    ('Invoice Import (CSV/JSON/UBL)', 'DONE',    '9/10', 'invoice-import.ts'),
    ('Invoice Clone',                 'PARTIAL', '4/10', 'Local state only, not persisted'),
    ('Invoice Locking After Sent',    'MISSING', '0/10', 'No readOnly mode in editor'),
    ('Invoice Attachments',           'MISSING', '0/10', 'Type exists, no UI'),
    ('Invoice Status Transitions',    'MISSING', '0/10', 'No validation in backend'),
    ('Business Letter CRUD',          'DONE',    '8/10', 'Shares invoice flow'),
    ('Letter Body Rich-Text',         'PARTIAL', '2/10', 'invoice.body typed, no UI rendered'),
    ('Letter Auto-Numbering',         'PARTIAL', '3/10', 'Format field exists, not wired'),
    ('Template CRUD',                 'DONE',    '9/10', 'TemplateLibrary + TemplateEditor'),
    ('Template Designer Drag-Drop',   'PARTIAL', '5/10', 'Handlers exist, no visual feedback'),
    ('Template Live Preview Cards',   'MISSING', '0/10', 'Cards show text only'),
    ('Buyer Management',              'DONE',    '9/10', 'BuyerController full CRUD'),
    ('Company Settings',              'PARTIAL', '6/10', 'Missing logo upload, format builder'),
    ('Logo Upload in Settings',       'PARTIAL', '3/10', 'URL input only, no file picker'),
    ('Number Format Builder',         'PARTIAL', '2/10', 'Plain input, no preview'),
    ('Signature Upload',              'MISSING', '0/10', 'Not in any settings screen'),
    ('Dashboard Statistics',          'DONE',    '8/10', 'Charts, recent invoices, actions'),
    ('Batch Invoice Validation',      'DONE',    '9/10', 'Dashboard bulk validate'),
    ('Workspace File Manager',        'DONE',    '8/10', 'Full CRUD + AI search'),
    ('Workspace in Navigation',       'MISSING', '0/10', 'No sidebar entry'),
    ('Workspace AI Search',           'PARTIAL', '6/10', 'Works but SQL injection risk'),
    ('Multi-language (4 langs)',      'DONE',    '9/10', 'EN/DE/AR/PL complete'),
    ('Authentication (login/signup)', 'DONE',    '9/10', 'JWT + Zustand'),
    ('Password Reset Flow',           'DONE',    '8/10', 'Email link + form'),
    ('Two-Factor Auth (TOTP)',        'MISSING', '0/10', 'Not implemented'),
    ('Customer Billing Screen',       'DONE',    '7/10', 'Plans, usage, upgrade works'),
    ('Payment History',               'PARTIAL', '1/10', 'Empty [] stub in backend'),
    ('Stripe Upgrade',                'DONE',    '7/10', 'Creates checkout session'),
    ('Stripe Webhook Sync',           'MISSING', '0/10', 'No handler, no DB table'),
    ('Plan Usage Enforcement',        'MISSING', '0/10', 'Zero enforcement in controllers'),
    ('Monthly Usage Reset',           'MISSING', '0/10', 'No cron job'),
    ('Quick Access Portal',           'DONE',    '7/10', 'OTP flow, draft persist'),
    ('Quick Access Rate Limiting',    'MISSING', '0/10', 'No OTP rate limit'),
    ('Activity / Audit Log',          'DONE',    '7/10', 'logAction() used, no field diffs'),
    ('Legal / CMS Pages',             'DONE',    '8/10', 'All 4 pages seeded and served'),
    ('Admin Login',                   'DONE',    '9/10', 'Separate admin JWT'),
    ('Admin Packages CRUD',           'DONE',    '7/10', 'Inline editing works'),
    ('Admin Package Services',        'DONE',    '8/10', 'Full CRUD'),
    ('Admin Upgrade Plan (stub)',      'MISSING', '0/10', 'Returns success, does nothing'),
    ('Admin Users List/View',         'DONE',    '7/10', 'With usage stats'),
    ('Admin Suspend/Activate',        'DONE',    '8/10', 'Works'),
    ('Admin Password Reset',          'PARTIAL', '2/10', 'Hardcoded "password123"'),
    ('Admin Billing List',            'PARTIAL', '5/10', 'Uses subscriptions as invoices'),
    ('Admin Invoice PDF',             'PARTIAL', '1/10', 'Returns mock PDF bytes'),
    ('Admin Revenue Dashboard',       'PARTIAL', '5/10', 'Estimated, not actual'),
    ('Admin Analytics Dashboard',     'PARTIAL', '5/10', 'Approximations, rand() filler'),
    ('Admin Analytics CSV Export',    'PARTIAL', '1/10', 'Hardcoded 2-row mock'),
    ('Admin Tickets',                 'DONE',    '6/10', 'CRUD + status + tracking'),
    ('Admin Wiki (read)',              'DONE',    '7/10', 'Full tree viewer'),
    ('Admin Wiki (edit)',              'MISSING', '0/10', 'Never built — started and stopped'),
    ('Admin CMS Page Editor',         'PARTIAL', '5/10', 'Endpoint exists, UI limited'),
    ('Admin RBAC',                    'PARTIAL', '2/10', 'Infrastructure built, never enforced'),
    ('Customer Portal',               'DONE',    '7/10', 'Dashboard + invoice view'),
    ('AI Invoice Creation',           'PARTIAL', '4/10', 'Backend done, no UI entry point'),
    ('AI Workspace Search',           'PARTIAL', '6/10', 'Works but SQL injection risk'),
    ('Email Notifications',           'MISSING', '0/10', 'No email service'),
    ('Batch PDF ZIP Export',          'PARTIAL', '3/10', 'UI exists, exportInvoicesBulk() stub'),
    ('PDF/A-3 with UBL embed',        'MISSING', '0/10', 'Type hinted, never implemented'),
    ('Invoice Sharing Link',          'PARTIAL', '2/10', 'Token system started, no endpoint'),
    ('OVERALL',                       'PARTIAL', '6.5/10','Production-ready core; ~25 gaps'),
]
tbl(
    ['Feature','Status','Score','Notes'],
    master,
    widths=[2.3,0.8,0.65,2.7]
)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
out = '/home/sivaji/Downloads/BillingTool/BillingTool_Complete_Analysis.docx'
doc.save(out)
print('Saved:', out)
