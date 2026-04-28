from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Colour palette ────────────────────────────────────────────────────────────
PURPLE     = RGBColor(0x6B, 0x21, 0xA8)   # heading purple
DARK       = RGBColor(0x1F, 0x2A, 0x3C)   # body dark
GREY       = RGBColor(0x6B, 0x72, 0x80)   # subtitle / caption
GREEN      = RGBColor(0x06, 0x6B, 0x2B)   # DONE
AMBER      = RGBColor(0x92, 0x40, 0x0E)   # PARTIAL
RED        = RGBColor(0x99, 0x17, 0x17)   # MISSING / HIGH
TBL_HEAD   = RGBColor(0x6B, 0x21, 0xA8)   # table header bg
TBL_ALT    = RGBColor(0xF5, 0xF3, 0xFF)   # table alt row

def set_cell_bg(cell, hex_color: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='D1D5DB', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(text, level=1, colour=PURPLE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = colour
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return p

def body(text, bold=False, colour=DARK, size=10.5, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    run.font.bold = bold
    run.font.color.rgb = colour
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, colour=DARK, size=10.5, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    run.font.color.rgb = colour
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p

def subbullet(text, colour=GREY, size=10):
    return bullet(text, colour=colour, size=size, indent=2)

def status_badge(text):
    if 'DONE' in text:        return (text, GREEN)
    if 'PARTIAL' in text:     return (text, AMBER)
    if 'MISSING' in text:     return (text, RED)
    if 'HIGH' in text:        return (text, RED)
    if 'MEDIUM' in text:      return (text, AMBER)
    if 'LOW' in text:         return (text, GREEN)
    return (text, DARK)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '6B21A8')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = 'F5F3FF' if ri % 2 == 1 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            txt, colour = status_badge(str(cell_text))
            run = p.add_run(txt)
            run.font.color.rgb = colour
            run.font.size = Pt(9.5)
    # column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '6B21A8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
run = p.add_run('BillingTool')
run.font.bold = True
run.font.size = Pt(32)
run.font.color.rgb = PURPLE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Implementation & Gap Analysis')
run2.font.size = Pt(18)
run2.font.color.rgb = DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Generated: 2026-04-23')
run3.font.size = Pt(11)
run3.font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading('1. Executive Summary')
divider()
body(
    'The BillingTool is a React + TypeScript SPA with a CodeIgniter 4 PHP backend. '
    'It is approximately 72% production-ready. Core invoicing, buyer management, '
    'authentication, multi-language support (EN/DE/AR/PL), and PDF export are fully '
    'functional. The main gaps are in the Template Designer (no real drag-drop), '
    'Admin Panel (screens partially stubbed), Business Letter UX, and enterprise-grade '
    'features (OAuth, 2FA, email notifications, RBAC).'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. FEATURE COMPLETENESS MATRIX
# ══════════════════════════════════════════════════════════════════════════════
heading('2. Feature Completeness Matrix')
divider()

matrix_rows = [
    ('Invoice CRUD',                      '10/10', 'DONE'),
    ('Invoice PDF Export',                '10/10', 'DONE'),
    ('Invoice EN 16931 Validation',       '10/10', 'DONE'),
    ('Buyer Management',                  '9/10',  'DONE'),
    ('Multi-language (EN/DE/AR/PL)',       '9/10',  'DONE'),
    ('Frontend Routing',                  '9/10',  'DONE'),
    ('Dashboard Statistics',              '8/10',  'DONE'),
    ('Backend API Layer',                 '8/10',  'DONE'),
    ('Database / Migrations',             '9/10',  'DONE'),
    ('CMS / Legal Pages',                 '8/10',  'DONE'),
    ('Toasts / Notifications',            '8/10',  'DONE'),
    ('TypeScript Types',                  '8/10',  'DONE'),
    ('Authentication (basic)',            '7/10',  'PARTIAL'),
    ('Invoice UBL XML Export',            '7/10',  'PARTIAL'),
    ('Business Letter Module',            '6/10',  'PARTIAL'),
    ('Template System',                   '7/10',  'PARTIAL'),
    ('Template Designer (drag-drop)',     '3/10',  'PARTIAL'),
    ('Company Settings',                  '6/10',  'PARTIAL'),
    ('Admin Panel',                       '5/10',  'PARTIAL'),
    ('Tickets / Wiki',                    '5/10',  'PARTIAL'),
    ('AI Assistant',                      '6/10',  'PARTIAL'),
    ('Quick Access Portal',               '5/10',  'PARTIAL'),
    ('OVERALL',                           '7.2/10','PARTIAL'),
]
add_table(
    ['Module', 'Score', 'Status'],
    matrix_rows,
    col_widths=[3.2, 0.9, 1.4]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. MODULE DETAILS
# ══════════════════════════════════════════════════════════════════════════════
heading('3. Module-by-Module Analysis')
divider()

# ── 3.1 Invoice ───────────────────────────────────────────────────────────────
heading('3.1  Invoice Module', level=2)
body('Status: DONE (8/10)', bold=True, colour=GREEN)
body('Implemented:', bold=True)
for item in [
    'Full CRUD: list, create, edit, delete',
    'Line items with dynamic add/remove and tax calculation  (src/utils/invoice-calculations.ts)',
    'EN 16931 validation  (src/utils/invoice-validation.ts)',
    'Status workflow: draft → validated → sent → paid / cancelled',
    'Client-side PDF export  (src/utils/invoice-pdf.ts)',
    'UBL XML, JSON, CSV export',
    'Search, filter, sort in list view  (src/components/screens/InvoiceList.tsx)',
    'Template selection per invoice',
    'Batch validation from Dashboard',
    'Import from file',
]:
    bullet(item)

body('Gaps:', bold=True)
for item in [
    'No status transition enforcement — any status can jump to any other in the backend controller',
    'No invoice locking after "sent" state — edits still allowed',
    'invoice.signed field exists in types but no signature capture UI',
    'invoice.attachments array exists in types but no file upload UI',
    'No invoice clone/copy functionality',
    'No bulk PDF export (ZIP of multiple invoices)',
    'AuditLog table exists but is not linked to invoice status changes',
]:
    bullet(item, colour=AMBER)

# ── 3.2 Business Letter ───────────────────────────────────────────────────────
heading('3.2  Business Letter Module', level=2)
body('Status: PARTIAL (6/10)', bold=True, colour=AMBER)
body('Implemented:', bold=True)
for item in [
    'templateType: "business_letter" in types  (src/types/invoice.ts:114)',
    'Separate DEFAULT_LETTER_LAYOUT  (src/utils/invoice-templates-defaults.ts)',
    'handleNewBusinessLetter() in App.tsx',
    'letterNumberFormat field in company profile',
    'InvoiceEditor hides currency/tax fields for letters',
    'Subject field replaces Due Date',
    'PDF renders description element as letter body',
    'List view filters by type',
]:
    bullet(item)

body('Gaps:', bold=True)
for item in [
    'No rich text editor for letter body (plain input only, using invoice.note)',
    'No letter-specific fields: salutation, closing, reference number',
    'No pre-built letter type templates (inquiry, quotation, complaint)',
    'Letter preview title still shows "Invoice Preview" in places',
    'Auto-numbering for letters not wired to letterNumberFormat from profile',
    'No letterhead / address-block positioning options',
]:
    bullet(item, colour=AMBER)

# ── 3.3 Template System ───────────────────────────────────────────────────────
heading('3.3  Template System', level=2)
body('Status: PARTIAL (7/10)', bold=True, colour=AMBER)
body('Implemented:', bold=True)
for item in [
    'Full template CRUD via TemplateLibrary + TemplateEditor',
    'Tab filter: Invoice vs. Business Letter',
    'Template type selector in create/edit form',
    'Platform default templates bundled (invoice + letter)',
    'Apply template → creates new invoice/letter pre-filled',
    'Logo, header, footer, colours, tax/payment defaults per template',
    'Template delete with confirmation warning',
]:
    bullet(item)

body('Gaps (HIGH IMPACT):', bold=True)
for item in [
    'Template Designer is read-only — TemplateDesignLayout.tsx renders the canvas but drag-drop is non-functional',
    'No live invoice preview in template library cards',
    'No per-element style controls (font size, colour, bold) in designer',
    'No template versioning or history',
    'No template export/import as JSON',
    'Template sharing between users not supported',
]:
    bullet(item, colour=AMBER)

# ── 3.4 Authentication ────────────────────────────────────────────────────────
heading('3.4  Authentication', level=2)
body('Status: PARTIAL (7/10)', bold=True, colour=AMBER)
body('Implemented:', bold=True)
for item in [
    'Login, logout, signup, token refresh  (src/stores/authStore.ts)',
    'JWT stored in localStorage via Zustand persist',
    'Auto-logout on 401  (src/services/api.ts:41–55)',
    'Forgot password and reset password endpoints',
    'Separate admin auth  (src/services/adminApi.ts)',
    'Protected routes via ProtectedRoute.tsx',
]:
    bullet(item)

body('Gaps:', bold=True)
for item in [
    'No OAuth / SSO (Google, GitHub, Microsoft)',
    'No Two-Factor Authentication (2FA / TOTP)',
    'customerApi.ts passes tokens manually per call — no interceptor, 401 refresh does not work for customer portal',
    'No session timeout warning',
    'No password strength rules in UI',
    'No email verification on signup',
]:
    bullet(item, colour=AMBER)

# ── 3.5 Settings ──────────────────────────────────────────────────────────────
heading('3.5  Company Settings', level=2)
body('Status: PARTIAL (6/10)', bold=True, colour=AMBER)
body('Implemented:', bold=True)
for item in [
    'Company profile: name, VAT ID, address, phone, email, banking (IBAN/BIC)',
    'Invoice defaults: currency, tax rate, payment terms',
    'invoiceNumberFormat and letterNumberFormat fields in profile model',
    'defaultTemplateId stored in profile',
]:
    bullet(item)

body('Gaps:', bold=True)
for item in [
    'Company logo upload missing from Settings screen — exists only in TemplateEditor',
    'Invoice number format builder — DB field present but no UI to configure the pattern',
    'Signature image upload not exposed in any UI screen',
    'No email notification preferences',
    'No API key management for customers (admin-only)',
    'No data export / account backup',
]:
    bullet(item, colour=AMBER)

# ── 3.6 Admin Panel ───────────────────────────────────────────────────────────
heading('3.6  Admin Panel', level=2)
body('Status: PARTIAL (5/10)', bold=True, colour=AMBER)
body('Screens implemented:', bold=True)
for item in [
    'SA Login, Dashboard, Packages, Package Services',
    'Users, User Details, Billing, Usage',
    'Settings, Invoice Form, Tickets, Ticket Details',
    'Wiki (read-only), CMS Pages',
]:
    bullet(item)

body('All backend routes present in api/app/Config/Routes.php (lines 194–261).', bold=False)

body('Gaps:', bold=True)
for item in [
    'No RBAC for admins — all authenticated admins see all screens equally',
    'Admin portal has no auth guard before rendering screens',
    'SADashboard charts may not fully populate all analytics fields',
    'Wiki is read-only — no admin edit UI despite updatePage endpoint existing',
    'SAPages (CMS editor) may be partially implemented',
    'No audit log UI for admin actions',
    'No tenant management screen',
    'Email template management missing',
    'Demo data seeding endpoints exist in backend but no admin UI trigger',
]:
    bullet(item, colour=AMBER)

# ── 3.7 PDF Generation ────────────────────────────────────────────────────────
heading('3.7  PDF Generation', level=2)
body('Status: DONE (8/10)', bold=True, colour=GREEN)
body('Implemented:', bold=True)
for item in [
    'Full client-side PDF via src/utils/invoice-pdf.ts',
    'Dynamic layout from template elements',
    'GiroCode / QR payment rendering',
    'Tax summary table, line items table, multi-page support',
    'Header/footer HTML rendering',
    'Business letter body block (el.type === "description")',
]:
    bullet(item)

body('Gaps:', bold=True)
for item in [
    'No server-side PDF generation — all client-side, blocks UI thread on large invoices',
    'No batch PDF export (multiple invoices as ZIP)',
    'PDF/A-3 with embedded UBL XML listed in types but not implemented',
    'Custom fonts not supported (hardcoded Helvetica)',
    'Signature image not rendered (only a placeholder line)',
]:
    bullet(item, colour=AMBER)

# ── 3.8 Translations ──────────────────────────────────────────────────────────
heading('3.8  Translations (EN / DE / AR / PL)', level=2)
body('Status: DONE (9/10)', bold=True, colour=GREEN)
for item in [
    'EN / DE / AR / PL — all ~1326–1333 lines per file',
    'All major screens covered with full key parity',
    'RTL support for Arabic via isRtl in useLanguage hook',
    'LanguageSwitcher with flag + language code display',
    'Automatic fallback to English for any missing key',
]:
    bullet(item)

body('Gaps:', bold=True)
for item in [
    'No locale-specific number formatting (DE uses comma decimal separator)',
    'No plural form rules per language',
    'Some newer dynamic strings may be hardcoded English in components',
    'No translation management UI — changes require code edits',
]:
    bullet(item, colour=AMBER)

# ── 3.9 Buyers ────────────────────────────────────────────────────────────────
heading('3.9  Buyer Management', level=2)
body('Status: DONE (9/10)', bold=True, colour=GREEN)
for item in [
    'Full CRUD: list, create, edit, delete with confirmation',
    'Search, filter, sort and pagination',
    'Autocomplete in InvoiceEditor (BuyerAutocomplete.tsx)',
    'Address stored as JSON (Migration 2026-02-16)',
]:
    bullet(item)
body('Minor gaps: no bulk CSV import, no buyer deactivation, no contact persons.', colour=GREY)

# ── 3.10 Dashboard ────────────────────────────────────────────────────────────
heading('3.10  Dashboard', level=2)
body('Status: DONE (8/10)', bold=True, colour=GREEN)
for item in [
    'Revenue statistics: total, paid, pending, draft',
    'Status distribution pie chart (recharts)',
    'Monthly trend line chart (last 6 months)',
    'Recent invoices list and quick actions',
    'Batch validation and import dialog',
]:
    bullet(item)
body('Minor gaps: no custom date ranges, no forecasting, no top-buyer insights.', colour=GREY)

# ══════════════════════════════════════════════════════════════════════════════
# 4. KNOWN BUGS
# ══════════════════════════════════════════════════════════════════════════════
heading('4. Known Bugs & Breaking Points')
divider()

bugs = [
    ('HIGH',   'AdminLayout.tsx',          'No per-role access control — all admins see all screens'),
    ('HIGH',   'InvoiceController.php',    'No state-transition validation — any status can be set to any other'),
    ('MEDIUM', 'customerApi.ts',           'No Axios interceptor — 401 token refresh does not trigger for customer portal'),
    ('MEDIUM', 'Backend controllers',      'No email sent on invoice status change (sent, paid)'),
    ('MEDIUM', 'invoice-pdf.ts',           'PDF generation is 100% client-side — no server fallback for heavy workloads'),
    ('LOW',    'TemplateDesignLayout.tsx', 'Designer canvas renders but drag-drop interaction is non-functional'),
    ('LOW',    'API responses',            'Inconsistent envelope: some return { data: T }, others return T directly'),
]
add_table(
    ['Severity', 'File', 'Issue'],
    bugs,
    col_widths=[0.85, 2.1, 3.55]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PRIORITISED BACKLOG
# ══════════════════════════════════════════════════════════════════════════════
heading('5. Prioritised Backlog')
divider()

heading('High Priority', level=2)
high = [
    ('1', 'Template Designer drag-drop',           'TemplateDesignLayout.tsx needs interactive element repositioning'),
    ('2', 'Invoice status transition enforcement', 'Backend should validate allowed transitions per workflow'),
    ('3', 'Company logo upload in Settings',       'Currently only accessible inside TemplateEditor'),
    ('4', 'Invoice number format builder UI',      'DB field exists (invoiceNumberFormat) but no configuration UI'),
    ('5', 'Admin RBAC',                            'Role/permission checks before rendering admin screens'),
]
add_table(['#', 'Task', 'Notes'], high, col_widths=[0.3, 2.2, 4.0])

heading('Medium Priority', level=2)
medium = [
    ('6',  'Two-Factor Authentication (2FA)',   'TOTP or SMS-based second factor'),
    ('7',  'Invoice locking after "sent"',      'Prevent edits after status transitions beyond draft'),
    ('8',  'Email notifications',              'Send emails on invoice status changes (sent/paid)'),
    ('9',  'Batch PDF export',                 'Download multiple invoices as a ZIP archive'),
    ('10', 'Letter body rich-text editor',     'Replace plain input with RichTextEditor for letter content'),
    ('11', 'Wiki admin edit UI',               'updatePage endpoint exists; no frontend form'),
    ('12', 'customerApi.ts interceptor',       'Add Axios interceptor for automatic 401 token refresh'),
]
add_table(['#', 'Task', 'Notes'], medium, col_widths=[0.3, 2.2, 4.0])

heading('Low Priority', level=2)
low = [
    ('13', 'OAuth / SSO',                  'Google, Microsoft login support'),
    ('14', 'PDF/A-3 with embedded UBL',    'Compliance requirement for some EU markets'),
    ('15', 'Invoice clone/copy',           'Duplicate an existing invoice with a new number'),
    ('16', 'Bulk operations',              'Select multiple invoices for batch export/delete'),
    ('17', 'Template versioning',          'History of changes per template'),
    ('18', 'Advanced dashboard analytics', 'Custom date ranges, forecasting, top-buyer insights'),
    ('19', 'Signature image upload',       'Capture and render actual signature in PDF'),
    ('20', 'Attachment support',           'File uploads linked to individual invoices'),
    ('21', 'Translation management UI',    'Edit translations without code changes'),
    ('22', 'Locale number formatting',     'DE comma decimal separator, PL currency rules'),
]
add_table(['#', 'Task', 'Notes'], low, col_widths=[0.3, 2.2, 4.0])

# ══════════════════════════════════════════════════════════════════════════════
# 6. BACKEND COVERAGE
# ══════════════════════════════════════════════════════════════════════════════
heading('6. Backend Controller Coverage')
divider()

controllers = [
    ('Auth',                   'DONE',    'login, signup, logout, refresh, me, forgotPassword, resetPassword'),
    ('InvoiceController',      'DONE',    'CRUD + index with filters'),
    ('InvoiceTemplateCtrl',    'DONE',    'CRUD'),
    ('BuyerController',        'DONE',    'CRUD'),
    ('CompanyProfileCtrl',     'DONE',    'index, update'),
    ('AuditLogController',     'DONE',    'index, show'),
    ('AIInvoiceController',    'DONE',    'parseInvoice'),
    ('WorkspaceController',    'DONE',    'list, upload, mkdir, delete, rename, download, aiSearch'),
    ('AdminAuth',              'DONE',    'login, logout, me, refresh'),
    ('AdminPackages',          'DONE',    'CRUD'),
    ('AdminUsers',             'DONE',    'index, show, suspend, activate, resetPassword, export'),
    ('AdminBilling',           'DONE',    'index, create, show, downloadPdf, revenue'),
    ('AdminAnalytics',         'DONE',    'dashboard, tenantUsage, usage, exportUsage'),
    ('AdminSettings',          'DONE',    'profile, password, API keys, system settings'),
    ('TicketController',       'DONE',    'create, index, update, tracking'),
    ('CmsController',          'DONE',    'getPage (public), listPages, updatePage (admin)'),
    ('AdminWiki',              'PARTIAL', 'Read-only — no write/edit endpoint exposed'),
    ('Customer',               'PARTIAL', 'dashboard, invoices, subscription, updateProfile, usage'),
    ('PDFController',          'MISSING', 'No server-side PDF generation — client-only'),
    ('EmailController',        'MISSING', 'No transactional email sending on events'),
]
add_table(
    ['Controller', 'Status', 'Notes'],
    controllers,
    col_widths=[1.9, 0.85, 3.75]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. TRANSLATION COVERAGE
# ══════════════════════════════════════════════════════════════════════════════
heading('7. Translation Coverage')
divider()

trans = [
    ('English (en)', '1326 lines', 'DONE',    'Reference language — complete'),
    ('German (de)',  '1326 lines', 'DONE',    'Full parity with EN'),
    ('Arabic (ar)',  '1326 lines', 'DONE',    'Full parity + RTL support'),
    ('Polish (pl)',  '1333 lines', 'DONE',    'Full parity — newly added'),
]
add_table(
    ['Language', 'File Size', 'Status', 'Notes'],
    trans,
    col_widths=[1.6, 1.1, 0.85, 3.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. DATABASE MIGRATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading('8. Database Migrations')
divider()

migs = [
    ('2020-01-15', 'InitialSchema',             'DONE', 'Core tables: users, tenants, invoices, templates'),
    ('2026-02-16', 'Buyers table',              'DONE', 'Buyer CRUD with address_json'),
    ('2026-02-20', 'WorkspaceFiles + AI',       'DONE', 'Workspace file storage + AI query history'),
    ('2026-02-27', 'QuickAccessSessions',       'DONE', 'OTP-based quick access'),
    ('2026-03-09', 'TicketTracking',            'DONE', 'Support ticket status tracking'),
    ('2026-03-10', 'UsageNotifications',        'DONE', 'Plan usage alerts'),
    ('2026-03-25', 'PackageServices',           'DONE', 'Service tier features'),
    ('2026-03-30', 'TenantUsage',               'DONE', 'Usage analytics per tenant'),
    ('2026-04-03', 'DefaultTemplate to Profile','DONE', 'defaultTemplateId on company profile'),
    ('2026-04-06', 'InvoiceNumberFormat',       'DONE', 'Number format + invoice defaults on profile'),
    ('2026-04-18', 'CmsPages + Seeds',          'DONE', 'CMS pages, Privacy/Terms/Cookie/Impressum seed'),
]
add_table(
    ['Migration Date', 'Description', 'Status', 'Notes'],
    migs,
    col_widths=[1.4, 2.0, 0.75, 2.35]
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. OVERALL SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading('9. Overall Summary')
divider()

summary = [
    ('Frontend Routing',   '9/10', 'DONE',    'No permission middleware'),
    ('Authentication',     '7/10', 'PARTIAL', 'Missing OAuth, 2FA, SSO'),
    ('Invoice Module',     '8/10', 'DONE',    'No locking, signing incomplete'),
    ('Business Letters',   '6/10', 'PARTIAL', 'Limited UX and styling'),
    ('Templates',          '7/10', 'PARTIAL', 'No interactive designer'),
    ('Buyers',             '9/10', 'DONE',    'Complete CRUD'),
    ('Settings',           '6/10', 'PARTIAL', 'Missing logo upload, format builder'),
    ('Admin Panel',        '5/10', 'PARTIAL', 'Many screens incomplete'),
    ('Translations',       '9/10', 'DONE',    'All four languages complete'),
    ('API Layer',          '8/10', 'DONE',    'Response format inconsistency'),
    ('Backend Controllers','8/10', 'DONE',    'Missing email / webhooks'),
    ('Migrations',         '9/10', 'DONE',    'All tables present'),
    ('PDF Generation',     '8/10', 'DONE',    'Client-side only'),
    ('CMS Pages',          '8/10', 'DONE',    'Limited admin editor UI'),
    ('Dashboard',          '8/10', 'DONE',    'Missing forecasting'),
    ('Notifications',      '8/10', 'DONE',    'Toast only, no email'),
    ('TypeScript Types',   '8/10', 'DONE',    'Minor gaps'),
    ('OVERALL',           '7.2/10','PARTIAL', 'Admin, Designer, Enterprise features'),
]
add_table(
    ['Area', 'Score', 'Status', 'Key Gap'],
    summary,
    col_widths=[1.9, 0.75, 0.95, 2.9]
)

body(
    'The two highest-impact gaps are: (1) the Template Designer — visually present but '
    'drag-drop is non-functional; (2) the Invoice Number Format UI — the database field '
    'and profile model exist but users have no way to configure the pattern through the UI.',
    size=10.5
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = '/home/sivaji/Downloads/BillingTool/BillingTool_Gap_Analysis.docx'
doc.save(out)
print(f'Saved: {out}')
